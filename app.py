"""
app.py
------
The Flask web application. Routes for:
  - landing/marketing page
  - signup, login, logout
  - forgot-password (security question flow)
  - the dashboard (profile, skills, documents, Cubic-Metric breakdown)
  - JSON API endpoints the dashboard's JavaScript calls to upload
    documents, add/delete skills, and re-run analysis — all WITHOUT a
    full page reload, which is what makes the Cubic-Metric bars feel
    "live."

HOW THE LIVE-UPDATE BEHAVIOR WORKS:
The dashboard page itself is rendered once, server-side, with whatever
data exists at page-load. After that, every action that could change
the Cubic-Metric numbers — uploading a document, adding a skill,
deleting a skill, re-running analysis — is a fetch() call from
static/js/dashboard.js to one of the /api/... routes below. Each of
those routes returns the SAME shape of JSON (see
pipeline.get_dashboard_state), and the JavaScript re-renders the
skills list and the dimension bars from whatever comes back. The
server is the single source of truth; the page never trusts its own
old state once a change has been made.
"""

import os
import time
import uuid
import json
from dotenv import load_dotenv
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, Response
import db
import auth
import pipeline
import analyzer
import extract
import identity
import cache as cache_module

load_dotenv()  # reads OPENAI_API_KEY and FLASK_SECRET_KEY from a local .env file

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-key-change-this-in-production")
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024  # 25MB upload cap per request

# One database connection is opened per request (on first use) and reused
# by every db.* call for the rest of that request, instead of each call
# opening its own -- this hook is what actually closes that shared
# connection once the response is done.
app.teardown_appcontext(db.close_db)


@app.after_request
def _cache_static_assets(response):
    # Static files (CSS/JS/icons) currently answer every single request
    # with Cache-Control: no-cache, forcing a network round-trip on every
    # page load just to get a 304. An hour of real caching lets repeat
    # loads skip the network entirely while still picking up a deploy
    # within an hour -- ETag-based revalidation (unchanged, still present)
    # keeps correctness beyond that window.
    if request.path.startswith("/static/"):
        if "v" in request.args:
            # Every template that links to this file appends ?v=<the
            # current deploy's asset_version> -- a deploy that changes the
            # file's content always changes that string too, so a URL
            # carrying one specific version can never start pointing at
            # different content later. That's exactly what max-age can be
            # unbounded (browser-capped at a year) for: no revalidation
            # request is ever needed, not even a 304 round-trip.
            response.headers["Cache-Control"] = "public, max-age=31536000, immutable"
        else:
            response.headers["Cache-Control"] = "public, max-age=3600"
    return response


# A stable-per-deploy string appended to every static CSS/JS URL as
# ?v=... so a fresh page load always gets the CSS/JS that matches the
# HTML it just rendered, even mid-way through the hour-long cache above
# -- otherwise a deploy that changes both a template and its stylesheet
# can serve new markup against a still-cached, now-mismatched
# stylesheet for up to an hour. Vercel sets VERCEL_GIT_COMMIT_SHA at
# build/runtime for every deployment; falling back to process start
# time covers local dev (still busts cache across restarts, just not
# within one).
_ASSET_VERSION = os.environ.get("VERCEL_GIT_COMMIT_SHA", "")[:10] or str(int(time.time()))


@app.context_processor
def _inject_asset_version():
    return {"asset_version": _ASSET_VERSION}

ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "txt", "jpg", "jpeg", "png", "tiff", "tif"}
ALLOWED_IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "gif", "webp"}
AVATAR_ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "avatars")


def _allowed_file(filename: str) -> bool:
    if not filename or not filename.strip() or "." not in filename:
        return False
    return filename.rsplit(".", 1)[-1].lower() in ALLOWED_EXTENSIONS


_db_initialized = False


@app.before_request
def _ensure_db():
    # init_db() is idempotent, but it still opens a connection and runs
    # a dozen-plus CREATE TABLE / ALTER TABLE statements — over a real
    # network connection to Postgres, doing that on every single
    # request (every page, every API call) adds real, avoidable latency
    # to everything the app does. The schema can't change mid-process,
    # so once is enough per running instance; a fresh serverless cold
    # start still gets its own fresh check via this same flag reset to
    # False when the process starts.
    global _db_initialized
    if not _db_initialized:
        db.init_db()
        _db_initialized = True


# Endpoints reachable even before a user has accepted the terms
# gate — everything else under login gets redirected to /terms until
# that's done. This check runs before the onboarding gate below, so a
# brand-new account always sees terms first, then onboarding.
_TERMS_EXEMPT_ENDPOINTS = {
    "terms_page",
    "api_disclaimer_accept",
    "api_delete_account",
    "logout_page",
    "static",
    "service_worker",
    "android_asset_links",
    "privacy_policy",
    "landing",
    "security_key_reveal",
    "api_security_key_acknowledge",
}


@app.before_request
def _require_terms():
    user = auth.current_user()
    if not user:
        return  # not logged in — normal login_required handling applies elsewhere
    if user.get("disclaimer_accepted"):
        return
    if request.endpoint in _TERMS_EXEMPT_ENDPOINTS:
        return
    return redirect(url_for("terms_page"))


# Endpoints reachable even before a user has confirmed their
# documents — everything else under login gets redirected to
# /onboarding until that's done. Static assets and public/auth pages
# are excluded too since they're not behind login_required anyway.
_ONBOARDING_EXEMPT_ENDPOINTS = {
    "onboarding_page",
    "api_onboarding_readout",
    "api_onboarding_upload",
    "api_onboarding_check",
    "api_onboarding_confirm",
    "api_onboarding_remove_document",
    "api_delete_account",
    "logout_page",
    "static",
    "service_worker",
    "android_asset_links",
    "privacy_policy",
    "terms_page",
    "api_disclaimer_accept",
    # The splash screen must always get a chance to render before any
    # further redirect happens — it hands the actual "where next"
    # decision to /dashboard itself, one hop later.
    "landing",
    "security_key_reveal",
    "api_security_key_acknowledge",
}


@app.before_request
def _require_onboarding():
    user = auth.current_user()
    if not user:
        return  # not logged in — normal login_required handling applies elsewhere
    if user.get("documents_confirmed"):
        return
    if request.endpoint in _ONBOARDING_EXEMPT_ENDPOINTS:
        return
    return redirect(url_for("onboarding_page"))


@app.context_processor
def inject_user():
    return {"current_user": auth.current_user()}


# ---------------- PUBLIC PAGES ----------------

@app.route("/sw.js")
def service_worker():
    # Served from the root (not /static/sw.js) so its default scope is
    # "/" instead of "/static/" — a service worker can only ever control
    # pages under the directory it's served from, and the PWA install
    # prompt (beforeinstallprompt) won't fire unless it actually
    # controls the page the user is on (/dashboard, etc).
    from flask import send_from_directory
    return send_from_directory(app.static_folder, "sw.js", mimetype="application/javascript")


@app.route("/.well-known/assetlinks.json")
def android_asset_links():
    # Lets Android's Digital Asset Links verifier confirm the Android
    # TWA app (signed with the matching certificate) is allowed to act
    # as this site's "app" — without this, the TWA still works but
    # falls back to showing a browser-style URL bar instead of the
    # full-screen native feel. The fingerprint below isn't secret (it's
    # meant to be published exactly like this) — it's the SHA-256 of the
    # release signing certificate generated for android/, matching what
    # .github/workflows/build-android-apk.yml signs the APK with.
    return jsonify([
        {
            "relation": ["delegate_permission/common.handle_all_urls"],
            "target": {
                "namespace": "android_app",
                "package_name": "com.employable.app",
                "sha256_cert_fingerprints": [
                    "67:DB:89:8E:9A:BA:6C:28:BF:B5:74:D4:08:C1:16:04:A4:E1:C1:9C:84:8B:B2:F6:78:C4:2F:C1:8D:09:81:35"
                ],
            },
        }
    ])


_ANDROID_APK_RELEASE_URL = "https://github.com/dsgnrcoza/mobile-employable/releases/download/android-latest/employable.apk"


@app.route("/download/android")
def download_android():
    # Proxies the built APK through our own domain instead of pointing
    # the browser straight at a github.com URL. Clicking a raw GitHub
    # link from inside the installed TWA (or this site running as an
    # installed PWA) means navigating outside the app's verified
    # origin, which hands the whole flow off to an external browser tab
    # instead of just downloading in place -- from the user's
    # perspective that looks exactly like "it redirects to GitHub"
    # instead of downloading. Keeping the entire request on our own
    # domain avoids that hand-off.
    import urllib.request
    import urllib.error

    try:
        req = urllib.request.Request(_ANDROID_APK_RELEASE_URL, headers={"User-Agent": "Employable-App"})
        with urllib.request.urlopen(req, timeout=25) as resp:
            data = resp.read()
    except (urllib.error.URLError, urllib.error.HTTPError):
        return jsonify({"error": "Couldn't fetch the Android app right now. Please try again shortly."}), 502

    return Response(
        data,
        mimetype="application/vnd.android.package-archive",
        headers={"Content-Disposition": 'attachment; filename="employable.apk"'},
    )


@app.route("/")
def landing():
    # Decide and redirect immediately, server-side — no splash screen
    # waiting on a client-side timer/animation-end event to fire before
    # it acts. That deferred pattern could get stuck on a real device
    # (slow network, a resource that never finishes loading, etc.),
    # leaving the user staring at a loading screen that never proceeds.
    # /dashboard itself still redirects on to /terms or /onboarding when
    # needed, so this one destination naturally covers every case: not
    # logged in -> /login, logged in but not onboarded -> /onboarding,
    # fully set up -> /dashboard.
    if not auth.current_user():
        return redirect(url_for("login_page"))
    return redirect(url_for("dashboard"))


@app.route("/signup", methods=["GET", "POST"])
def signup_page():
    if auth.current_user():
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        try:
            _user_id, security_key = auth.signup(
                full_name=request.form.get("full_name", ""),
                email=request.form.get("email", ""),
                password=request.form.get("password", ""),
                confirm_password=request.form.get("confirm_password", ""),
            )
            # Stashed in the session until the reveal page's "Continue"
            # button acknowledges it (see api_security_key_acknowledge) --
            # the plaintext key is never written anywhere more durable
            # than this.
            session["_reveal_security_key"] = security_key
            return redirect(url_for("security_key_reveal"))
        except auth.AuthError as e:
            flash(str(e), "error")

    return render_template("signup.html")


@app.route("/security-key")
@auth.login_required
def security_key_reveal():
    """
    Shows a freshly generated security key. Reached right after signup,
    and after any regenerate/reset that issues a new key. Left in the
    session (not popped here) until the "Continue" button explicitly
    acknowledges it via /api/security-key/acknowledge -- signup's own
    fetch-then-celebrate flow already makes one GET of this page in the
    background to detect success, before the user ever sees or clicks
    anything, so popping on GET would clear it before the real,
    user-visible navigation that follows.
    """
    security_key = session.get("_reveal_security_key")
    if not security_key:
        return redirect(url_for("dashboard"))
    return render_template("security_key_reveal.html", security_key=security_key)


@app.route("/api/security-key/acknowledge", methods=["POST"])
@auth.login_required
def api_security_key_acknowledge():
    session.pop("_reveal_security_key", None)
    return jsonify({"ok": True})


@app.route("/login", methods=["GET", "POST"])
def login_page():
    if auth.current_user():
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        try:
            auth.login(request.form.get("email", ""), request.form.get("password", ""))
            return redirect(url_for("dashboard"))
        except auth.AuthError as e:
            flash(str(e), "error")

    return render_template("login.html")


@app.route("/logout")
def logout_page():
    auth.log_out_user()
    flash("You've been logged out.", "success")
    return redirect(url_for("login_page"))


@app.route("/forgot-password")
def forgot_password_page():
    """
    A single form (email + security key + new password), driven by
    static/js/forgot-password.js via the JSON endpoint below -- there's
    no separate "request a code" step since the security key itself is
    the proof of identity.
    """
    if auth.current_user():
        return redirect(url_for("dashboard"))
    return render_template("forgot_password.html")


@app.route("/api/reset-password", methods=["POST"])
def api_reset_password():
    if auth.current_user():
        return jsonify({"ok": False, "error": "Already signed in."}), 400
    data = request.get_json(force=True)
    try:
        new_security_key = auth.reset_password_with_security_key(
            data.get("email", ""),
            data.get("security_key", ""),
            data.get("new_password", ""),
            data.get("confirm_password", ""),
        )
        # reset_password_with_security_key already logs the user in --
        # reusing the same one-time reveal page as signup keeps there
        # being exactly one path that ever shows a plaintext key.
        session["_reveal_security_key"] = new_security_key
        return jsonify({"ok": True, "redirect": url_for("security_key_reveal")})
    except auth.AuthError as e:
        return jsonify({"ok": False, "error": str(e)}), 400


# ---------------- ONBOARDING (required before first dashboard view) ----------------

@app.route("/onboarding")
@auth.login_required
def onboarding_page():
    user = auth.current_user()
    if user.get("documents_confirmed"):
        return redirect(url_for("dashboard"))
    documents = db.get_documents_for_user(user["id"])
    cv_documents = [d for d in documents if d.get("category") == "cv"]
    supporting_documents = [d for d in documents if d.get("category") != "cv"]
    return render_template(
        "onboarding.html",
        cv_documents=cv_documents,
        supporting_documents=supporting_documents,
        has_cv=pipeline.any_document_looks_like_cv(user["id"]),
    )


@app.route("/api/onboarding/readout", methods=["POST"])
@auth.login_required
def api_onboarding_readout():
    """
    The onboarding "wow" moment: parses the CV(s) just uploaded and
    returns a short, honest readout -- years of experience, a skills
    count, and 1-2 concrete problems that would hurt in a real ATS/
    recruiter pass. Real analysis of this specific person's real CV,
    not canned copy.
    """
    from openai import OpenAI
    user = auth.current_user()
    documents = db.get_documents_for_user(user["id"])
    cv_documents = [d for d in documents if d.get("category") == "cv"]
    if not cv_documents:
        return jsonify({"ok": False, "error": "Upload a CV first."}), 400

    doc_texts = []
    for d in cv_documents:
        text = (d.get("content") or "").strip()
        if text:
            doc_texts.append(f"[{d['filename']}]\n{text[:6000]}")
    doc_content_block = "\n\n---\n\n".join(doc_texts) if doc_texts else "No readable content extracted."

    system = (
        "You are Ploy's CV parser -- sharp, direct, zero corporate filler. "
        "Given the real content of a CV below, return a JSON object with exactly these keys:\n"
        "- 'years_experience': a short string like '5+ years', '2 years', or 'Entry-level' -- your best "
        "honest estimate from the actual dates/roles in the CV, never invented.\n"
        "- 'skills_count': an integer count of distinct real skills you can identify in the CV.\n"
        "- 'flags': an array of 1-2 short, specific, concrete problems with THIS CV that would hurt it in "
        "a real ATS scan or recruiter skim (e.g. no measurable results, no dedicated skills section, dates "
        "that don't parse cleanly, a wall of text with no structure). Each flag is one short sentence, "
        "specific to what's actually in this document -- never generic advice. If the CV is genuinely solid "
        "with nothing significant to flag, return an empty array.\n"
        "Never invent facts not present in the CV."
    )
    prompt = f"CV content:\n{doc_content_block}"

    try:
        client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            max_tokens=500,
            response_format={"type": "json_object"},
        )
        parsed = json.loads(resp.choices[0].message.content.strip())
        return jsonify({
            "ok": True,
            "years_experience": parsed.get("years_experience", ""),
            "skills_count": int(parsed.get("skills_count") or 0),
            "flags": [f for f in (parsed.get("flags") or []) if f][:2],
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/onboarding/upload", methods=["POST"])
@auth.login_required
def api_onboarding_upload():
    user = auth.current_user()
    files = request.files.getlist("documents")
    if not files or all(f.filename == "" for f in files):
        return jsonify({"ok": False, "error": "No files were selected."}), 400

    category = (request.form.get("category") or "").strip()

    saved, rejected = [], []
    for f in files:
        if f.filename == "":
            continue
        if not _allowed_file(f.filename):
            rejected.append(f.filename)
            continue
        saved.append(pipeline.save_uploaded_file(user["id"], f, category=category))

    if not saved:
        return jsonify({"ok": False, "error": "None of the selected files are a supported type."}), 400

    result = pipeline.check_identity_conflict(user["id"])
    documents = db.get_documents_for_user(user["id"])
    return jsonify(
        {
            "ok": True,
            "rejected": rejected,
            "documents": [
                {
                    "id": d["id"],
                    "filename": d["filename"],
                    "file_type": d["file_type"],
                    "category": d.get("category", ""),
                    "file_size": d.get("file_size"),
                }
                for d in documents
            ],
            "has_cv": pipeline.any_document_looks_like_cv(user["id"]),
            **result,
        }
    )


@app.route("/api/onboarding/check")
@auth.login_required
def api_onboarding_check():
    user = auth.current_user()
    return jsonify({
        "has_cv": pipeline.any_document_looks_like_cv(user["id"]),
        **pipeline.check_identity_conflict(user["id"]),
    })


@app.route("/api/onboarding/document/<int:document_id>", methods=["DELETE"])
@auth.login_required
def api_onboarding_remove_document(document_id):
    user = auth.current_user()
    db.delete_document(user["id"], document_id)
    return jsonify({
        "ok": True,
        "has_cv": pipeline.any_document_looks_like_cv(user["id"]),
        **pipeline.check_identity_conflict(user["id"]),
    })


@app.route("/api/onboarding/confirm", methods=["POST"])
@auth.login_required
def api_onboarding_confirm():
    """
    Finalizes onboarding. The frontend sends the name the user
    confirmed (e.g. picked from the conflict chooser, or the single
    guessed name when there was no conflict) plus the list of
    document ids that belong to that person. Any document NOT in
    that list is deleted — mixing two people's documents in one
    account is never allowed, so the only way forward is to keep one
    identity's files and drop the rest.
    """
    data = request.json if request.is_json else request.form
    keep_ids = data.get("keep_document_ids") or []
    try:
        keep_ids = [int(i) for i in keep_ids]
    except (TypeError, ValueError):
        return jsonify({"ok": False, "error": "Invalid document selection."}), 400

    user = auth.current_user()
    documents = db.get_documents_for_user(user["id"])
    if not documents:
        return jsonify({"ok": False, "error": "Please upload at least one document first."}), 400

    # The name the user typed in onboarding's first step is the
    # source of truth now -- only fall back to a document-guessed name
    # (via the client-sent "name", populated from the identity-conflict
    # chooser) if that first step somehow never ran.
    chosen_name = (data.get("name") or "").strip() or (user.get("full_name") or "").strip()

    pipeline.resolve_identity_conflict(user["id"], keep_ids or [d["id"] for d in documents])

    remaining = db.get_documents_for_user(user["id"])
    if not remaining:
        return jsonify({"ok": False, "error": "No documents left after removing the conflicting set."}), 400

    db.set_documents_confirmed(user["id"], chosen_name)

    try:
        pipeline.run_analysis_for_user(user["id"])
    except analyzer.CVAnalyzerError as e:
        # Onboarding is still considered complete — scoring can be
        # re-run from the dashboard's "Refresh Score" button.
        return jsonify({"ok": True, "warning": f"Documents confirmed, but scoring couldn't run yet: {e}"})

    return jsonify({"ok": True})


# ---------------- DASHBOARD ----------------

@app.route("/dashboard")
@auth.login_required
def dashboard():
    # Ploy's chat screen (see home.html) -- kept at this same route/
    # endpoint name since every login/onboarding redirect in this file
    # targets url_for("dashboard"). The old dashboard.html template
    # (Cubic-Metric score, CV Workshop, etc.) is untouched on disk,
    # just no longer rendered from here.
    user = auth.current_user()
    profile = {
        "full_name": user.get("full_name") or "",
        "avatar_url": _avatar_url_for(user),
        "initials": _initials(user.get("full_name") or ""),
    }
    # Looked up here instead of left to the client's first fetch -- the
    # chat screen's initial paint (resuming the most recent conversation)
    # no longer waits on a round-trip after the page has already loaded.
    conversations = db.get_conversations_for_user(user["id"])
    initial_conversation_id = conversations[0]["id"] if conversations else None
    initial_messages = (
        _serialize_conversation_messages(initial_conversation_id, user["id"])
        if initial_conversation_id else None
    )
    return render_template(
        "home.html", profile=profile, conversations=conversations,
        initial_conversation_id=initial_conversation_id, initial_messages=initial_messages,
    )


def _avatar_url_for(user):
    avatar_path = user.get("avatar_path") or ""
    if avatar_path.startswith("data:"):
        return avatar_path
    if avatar_path:
        return f"/static/{avatar_path}"
    return ""


def _initials(full_name):
    parts = [p for p in full_name.strip().split() if p]
    if not parts:
        return "?"
    if len(parts) == 1:
        return parts[0][0].upper()
    return (parts[0][0] + parts[-1][0]).upper()


def _real_uploaded_docs(user_id):
    """Documents this user actually uploaded -- excludes CVs/letters Ploy
    itself generated, so every new generation is grounded in the user's
    original source material rather than compounding on its own earlier
    output."""
    return [d for d in db.get_documents_for_user(user_id) if d.get("category") not in ("generated_cv", "generated_letter")]


def _doc_content_block(docs, doc_char_cap=3000, total_char_cap=20000):
    doc_texts = []
    total_chars = 0
    for d in docs:
        if total_chars >= total_char_cap:
            break
        try:
            txt = (d.get("content") or "").strip()
            if not txt and d.get("stored_path") and os.path.exists(d["stored_path"]):
                txt = (extract.extract_text(d["stored_path"]) or "").strip()
            if txt:
                snippet = txt[:doc_char_cap]
                doc_texts.append(f"[{d['filename']}]\n{snippet}")
                total_chars += len(snippet)
        except Exception:
            pass
    return "\n\n---\n\n".join(doc_texts) if doc_texts else "No documents uploaded yet."


@app.route("/api/verdict", methods=["POST"])
@auth.login_required
def api_verdict():
    """Powers the Verdict card (chip 01 / "Am I a fit for this job?").
    Takes a raw pasted job ad, returns a structured fit verdict scored
    against this user's real documents -- rendered client-side as a card,
    never as a wall of chat text."""
    user = auth.current_user()
    data = request.get_json(force=True)
    job_ad = (data.get("job_ad") or "").strip()[:8000]
    if not job_ad:
        return jsonify({"ok": False, "error": "Paste a job ad first."}), 400
    try:
        return jsonify(_run_verdict(user, job_ad))
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


def _build_skills_chart_card(user, chart_type):
    """Builds a chart card from the user's real, already-computed
    Employability analysis (the same 8-dimension breakdown the old
    dashboard used) -- never fabricated, and returns None if they don't
    have one yet (caller replies with a plain-text prompt to run one
    instead of showing an empty chart)."""
    analysis = db.get_latest_analysis(user["id"])
    if not analysis:
        return None
    try:
        result = json.loads(analysis["result_json"])
    except (TypeError, ValueError, KeyError):
        return None
    dimensions = [d for d in (result.get("dimensions") or []) if d.get("label") and d.get("score") is not None]
    if not dimensions:
        return None
    return {
        "ok": True,
        "type": "chart",
        "chart_type": chart_type if chart_type in ("bar", "pie") else "bar",
        "labels": [d["label"] for d in dimensions],
        "values": [round(float(d["score"]), 1) for d in dimensions],
    }


def _generate_chat_image(prompt):
    """Image Generator plugin: generates via OpenAI's image model using
    the same API key/client already configured for chat -- no separate
    third-party service or credential to wire up. Returns base64 PNG
    data; raises on failure so the caller can turn it into a plain-text
    error reply."""
    from openai import OpenAI
    client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
    resp = client.images.generate(model="dall-e-3", prompt=prompt, size="1024x1024", n=1, response_format="b64_json")
    return resp.data[0].b64_json


def _run_verdict(user, job_ad):
    """Shared by /api/verdict and /api/chat's check_job_fit tool call --
    one grounded-scoring implementation, whether the user reached it by
    tapping chip 01 or just typing the equivalent request. Raises on
    failure; callers turn that into a JSON error response."""
    from openai import OpenAI

    doc_content_block = _doc_content_block(_real_uploaded_docs(user["id"]))

    system = (
        "You are Ploy's brutally honest fit evaluator. Given a pasted job ad and the real content of "
        "this user's uploaded documents, decide how strong a candidate they are for THIS specific job, "
        "right now -- not how they could look someday.\n\n"
        "GROUNDING: base the verdict only on what's actually in the documents below. Never invent "
        "experience, skills, or qualifications that aren't there. If there isn't enough real document "
        "content to judge, say so in 'breakdown' and score low.\n\n"
        "OUTPUT RULES -- return a JSON object with exactly these keys:\n"
        "- 'job_title': the job title, extracted from the ad.\n"
        "- 'company': the company name if stated in the ad, else an empty string.\n"
        "- 'location': the job location if stated in the ad, else an empty string.\n"
        "- 'fit_score': an integer 0-100. 0-49 means not currently a fit, 50-74 means a partial/borderline "
        "fit, 75-100 means a strong fit. Be honest, not encouraging -- a mediocre match belongs in the "
        "40s-60s, not inflated into the 80s.\n"
        "- 'strengths': an array of at most 3 short strings (max ~12 words each), specific things in their "
        "real documents that support this exact job.\n"
        "- 'gaps': an array of at most 3 short strings (max ~12 words each), specific things this exact "
        "job needs that their documents don't show. Empty array only if there is truly nothing missing.\n"
        "- 'breakdown': 2-4 sentences of fuller reasoning behind the score, citing specifics from their "
        "real documents and the job ad.\n\n"
        "No markdown formatting inside any field -- plain text only."
    )
    prompt = f"Job ad pasted by the user:\n{job_ad}\n\nFull content of the user's actual uploaded documents:\n{doc_content_block}"

    client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
        max_tokens=700,
        response_format={"type": "json_object"},
    )
    parsed = json.loads(resp.choices[0].message.content.strip())
    fit_score = max(0, min(100, int(parsed.get("fit_score") or 0)))
    job_title = (parsed.get("job_title") or "").strip()[:200] or "This role"
    company = (parsed.get("company") or "").strip()[:200]
    return {
        "ok": True,
        "type": "verdict",
        "job_title": job_title,
        "company": company,
        "location": (parsed.get("location") or "").strip()[:200],
        "fit_score": fit_score,
        "previous_fit_score": db.get_previous_fit_score(user["id"], job_title, company),
        "strengths": [s for s in (parsed.get("strengths") or []) if s][:3],
        "gaps": [s for s in (parsed.get("gaps") or []) if s][:3],
        "breakdown": (parsed.get("breakdown") or "").strip(),
        "job_ad": job_ad,
    }


def _run_gap_analysis(user, target_role):
    """Powers the Gap Analysis card -- the "What's holding me back?" move
    promised by the system prompt and the chat-gaps chip, scoring the
    user's real documents against what a given role/field typically
    needs and returning concrete, specific gaps rather than generic
    advice. Raises on failure; callers turn that into a JSON error
    response."""
    from openai import OpenAI

    doc_content_block = _doc_content_block(_real_uploaded_docs(user["id"]))

    system = (
        "You are Ploy's honest gap analyst. Given the real content of this user's uploaded documents and "
        "the role or field they're targeting, identify what's actually holding them back from being a "
        "strong candidate for that kind of role right now.\n\n"
        "GROUNDING: base this only on what's actually in the documents below versus what that role/field "
        "typically and genuinely requires. Never invent experience they don't have, and never invent a "
        "requirement that isn't a real, common expectation for that role.\n\n"
        "OUTPUT RULES -- return a JSON object with exactly these keys:\n"
        "- 'target_role': the role/field, cleaned up, as given.\n"
        "- 'readiness_score': an integer 0-100 for how ready they are for that kind of role right now. Be "
        "honest, not encouraging -- a real gap belongs in the 30s-60s, not inflated.\n"
        "- 'strengths': an array of at most 4 short strings (max ~12 words each), specific things in their "
        "real documents that already work in that direction.\n"
        "- 'gaps': an array of at most 5 short strings (max ~12 words each), specific, concrete things "
        "missing -- a skill, a certification, a type of experience, a portfolio -- not vague advice like "
        "'gain more experience'.\n"
        "- 'breakdown': 2-4 sentences of fuller reasoning, citing specifics from their real documents.\n\n"
        "No markdown formatting inside any field -- plain text only."
    )
    prompt = f"Role/field the user is targeting: {target_role}\n\nFull content of the user's actual uploaded documents:\n{doc_content_block}"

    client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
        max_tokens=700,
        response_format={"type": "json_object"},
    )
    parsed = json.loads(resp.choices[0].message.content.strip())
    readiness_score = max(0, min(100, int(parsed.get("readiness_score") or 0)))
    return {
        "ok": True,
        "type": "gap",
        "target_role": (parsed.get("target_role") or target_role or "").strip()[:200] or "This role",
        "readiness_score": readiness_score,
        "strengths": [s for s in (parsed.get("strengths") or []) if s][:4],
        "gaps": [s for s in (parsed.get("gaps") or []) if s][:5],
        "breakdown": (parsed.get("breakdown") or "").strip(),
    }


def _filename_stub(full_name):
    first = (full_name or "").strip().split()[:1]
    stub = "".join(ch for ch in (first[0] if first else "My") if ch.isalnum())
    return stub or "My"


def _company_stub(company):
    stub = "".join(ch for ch in (company or "") if ch.isalnum())
    return stub


def _generate_tailored_document(user, kind, job_title="", company="", job_ad="", template=None):
    """Shared by /api/document and /api/letter-document -- generates a
    fresh, ATS-safe CV or cover letter grounded in the user's real
    uploaded documents (never Ploy's own earlier output), tailored to a
    specific job when one is given, renders it to PDF, and stores it as
    a private per-user document. Returns a Document-card-shaped dict, or
    raises on failure (caller turns that into a JSON error response)."""
    from openai import OpenAI

    real_docs = _real_uploaded_docs(user["id"])
    doc_content_block = _doc_content_block(real_docs)
    job_context = (
        f"Tailor this specifically for the following job:\nTitle: {job_title or 'Not specified'}\n"
        f"Company: {company or 'Not specified'}\nJob ad:\n{job_ad}\n\n"
        if job_ad else
        "No specific job was given -- write a strong, general-purpose version tailored to the user's own real experience.\n\n"
    )

    if kind == "cv":
        doc_label = "CV"
        design_guidance = (
            "Apply real CV design sense: a clear name/contact header, short bold section headings "
            "(Experience, Education, Skills, etc.) in a consistent order, reverse-chronological entries, "
            "bullet points for achievements rather than dense paragraphs, generous spacing so it's scannable "
            "in a 6-second recruiter skim. Keep it ATS-friendly: no unusual layouts, plain section headings a "
            "parser would recognize.\n\n"
            f"{CV_STRUCTURE_GUIDANCE}"
        )
    else:
        doc_label = "cover letter"
        design_guidance = (
            "Write a genuine cover letter: a brief opening naming the role, 2-3 short paragraphs connecting "
            "the user's real, specific experience to what this job actually needs, a direct closing. No "
            "generic filler paragraphs ('I am writing to express my interest...'). Sharp and specific, not "
            "corporate boilerplate.\n\n"
            "STRUCTURE -- follow this exact shape (the app applies its own visual template on top):\n"
            "  <p class=\"letter-date\">today's real date</p>\n"
            "  <p class=\"letter-greeting\">Dear Hiring Manager,</p> (or a real named greeting if known)\n"
            "  <p>body paragraph</p> (2-3 of these)\n"
            "  <p class=\"letter-signoff\">Warm regards,<br>Full Name</p>"
        )

    system = (
        f"You are Ploy's document writer, generating a {doc_label} for this user.\n\n"
        f"{job_context}"
        f"{design_guidance}\n\n"
        "GROUNDING -- the most important rule, above all others: every fact (employer names, job titles, "
        "dates, schools, achievements, contact details) must come from the user's real documents below. "
        "NEVER invent a person, career, employer, or achievement that isn't actually theirs. If there isn't "
        "enough real information to write this, say so plainly in 'description' instead of fabricating.\n\n"
        "OUTPUT RULES -- return a JSON object with exactly these keys:\n"
        "- 'html': the full document as valid HTML using <p>, <strong>, <em>, <ul>, <li>, <h1>, <h2>, <h3>, "
        "<hr>, <br> tags"
        + (", plus the exact cv-header/cv-entry-* classes described above" if kind == "cv" else "")
        + ". No markdown, no code fences, no <html>/<head>/<body> wrappers.\n"
        "- 'description': one short sentence (max 20 words) describing what was produced.\n"
        + ("- 'fit_score': an integer 0-100 estimating how strong a fit THIS tailored document now makes for "
           "the job above, judged the same way a recruiter would score the original documents.\n" if job_ad else "") +
        f"\nThis user's name: {user.get('full_name') or 'Unknown'}\n\n"
        f"Full content of the user's actual uploaded documents (the ONLY source of truth):\n{doc_content_block}"
    )

    client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": system}, {"role": "user", "content": f"Generate the {doc_label} now."}],
        max_tokens=2400,
        response_format={"type": "json_object"},
    )
    parsed = json.loads(resp.choices[0].message.content.strip())
    html = parsed.get("html") or ""
    if not html.strip():
        raise ValueError(parsed.get("description") or "Not enough real document content to generate this.")

    pdf_bytes = _render_cv_pdf_bytes(cv_html=html, template=template if kind == "cv" else None)
    if pdf_bytes is None:
        raise ValueError("Couldn't render a document from that content.")

    name_stub = _filename_stub(user.get("full_name"))
    company_stub = _company_stub(company)
    label_stub = "CV" if kind == "cv" else "CoverLetter"
    filename = f"{name_stub}_{label_stub}_{company_stub}.pdf" if company_stub else f"{name_stub}_{label_stub}.pdf"

    import base64 as _b64
    document_id = db.add_document(
        user["id"], filename, stored_path="", file_type="pdf",
        content=html, category=f"generated_{kind}",
        file_size=len(pdf_bytes), file_bytes_b64=_b64.b64encode(pdf_bytes).decode("utf-8"),
    )

    fit_score = parsed.get("fit_score")
    if fit_score is not None:
        try:
            fit_score = max(0, min(100, int(fit_score)))
        except (TypeError, ValueError):
            fit_score = None

    return {
        "ok": True,
        "type": "document",
        "kind": kind,
        "document_id": document_id,
        "filename": filename,
        "job_title": job_title,
        "company": company,
        "fit_score": fit_score,
    }


@app.route("/api/document", methods=["POST"])
@auth.login_required
def api_document():
    """Powers the Document card for a tailored CV -- reached via chip 02
    ("Build my CV") or the Verdict card's "Fix my CV for this job" button."""
    user = auth.current_user()
    data = request.get_json(force=True)
    try:
        card = _generate_tailored_document(
            user, "cv",
            job_title=(data.get("job_title") or "").strip()[:200],
            company=(data.get("company") or "").strip()[:200],
            job_ad=(data.get("job_ad") or "").strip()[:8000],
            template=data.get("template") if data.get("template") in CV_TEMPLATES else None,
        )
        return jsonify(card)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/letter-document", methods=["POST"])
@auth.login_required
def api_letter_document():
    """Powers the Document card for a cover letter -- reached via the
    "Cover letter" button on an existing CV Document card."""
    user = auth.current_user()
    data = request.get_json(force=True)
    try:
        card = _generate_tailored_document(
            user, "letter",
            job_title=(data.get("job_title") or "").strip()[:200],
            company=(data.get("company") or "").strip()[:200],
            job_ad=(data.get("job_ad") or "").strip()[:8000],
        )
        return jsonify(card)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/document-download/<int:document_id>")
@auth.login_required
def api_document_download(document_id):
    """Serves a generated CV/cover-letter PDF -- scoped to the requesting
    user's own documents, same as every other per-user document access
    in this app."""
    from flask import send_file
    import io
    user = auth.current_user()
    doc_row = db.get_document_by_id(user["id"], document_id)
    file_bytes = db.get_document_file_bytes(user["id"], document_id)
    if not doc_row or file_bytes is None:
        return jsonify({"ok": False, "error": "Document not found."}), 404
    return send_file(io.BytesIO(file_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=doc_row.get("filename") or "document.pdf")


@app.route("/api/document/<int:document_id>/save", methods=["POST"])
@auth.login_required
def api_document_save(document_id):
    """Writes Builder edits back to a generated CV/cover-letter -- so a
    card's "Download" later reflects the edited version, not the
    original AI draft. Scoped to the requesting user's own documents."""
    import base64 as _b64
    user = auth.current_user()
    doc_row = db.get_document_by_id(user["id"], document_id)
    if not doc_row or doc_row.get("category") not in ("generated_cv", "generated_letter"):
        return jsonify({"ok": False, "error": "Document not found."}), 404

    data = request.get_json(force=True)
    html = (data.get("html") or "").strip()
    template = data.get("template") if (doc_row.get("category") == "generated_cv" and data.get("template") in CV_TEMPLATES) else None
    pdf_bytes = _render_cv_pdf_bytes(cv_html=html, template=template)
    if pdf_bytes is None:
        return jsonify({"ok": False, "error": "Nothing to save."}), 400

    db.update_document_content(user["id"], document_id, content=html,
                               file_bytes_b64=_b64.b64encode(pdf_bytes).decode("utf-8"),
                               file_size=len(pdf_bytes))
    return jsonify({"ok": True})


@app.route("/conversations")
@auth.login_required
def conversations_page():
    return render_template("conversations.html")


@app.route("/plugins")
@auth.login_required
def plugins_page():
    user = auth.current_user()
    try:
        enabled = json.loads(user.get("enabled_plugins") or "[]")
    except Exception:
        enabled = []
    return render_template("plugins.html", plugins=PLUGINS, enabled=enabled)


@app.route("/api/plugins", methods=["POST"])
@auth.login_required
def api_set_plugin_enabled():
    user = auth.current_user()
    data = request.get_json(force=True)
    key = (data.get("key") or "").strip()
    if key not in PLUGINS:
        return jsonify({"ok": False, "error": "Unknown plugin."}), 404
    try:
        enabled = set(json.loads(user.get("enabled_plugins") or "[]"))
    except Exception:
        enabled = set()
    if data.get("enabled"):
        enabled.add(key)
    else:
        enabled.discard(key)
    db.set_enabled_plugins(user["id"], sorted(enabled))
    return jsonify({"ok": True, "enabled": sorted(enabled)})


@app.route("/api/upload", methods=["POST"])
@auth.login_required
def api_upload():
    user = auth.current_user()
    files = request.files.getlist("documents")
    if not files or all(f.filename == "" for f in files):
        return jsonify({"ok": False, "error": "No files were selected."}), 400

    # Capture score before upload so we can compute per-document delta
    _pre_row = db.get_latest_analysis(user["id"])
    _pre_analysis = json.loads(_pre_row["result_json"]) if _pre_row else {}
    _score_before = float(_pre_analysis.get("overall_score") or 0)
    _dims_before = {d["label"]: d["score"] for d in (_pre_analysis.get("dimensions") or [])}

    saved = []
    rejected = []
    identity_rejected = []
    for f in files:
        if f.filename == "":
            continue
        if not _allowed_file(f.filename):
            rejected.append(f.filename)
            continue

        # Read the file once so we can check identity before deciding
        # whether to keep it, then rewind so save_uploaded_file can
        # still write the original bytes to disk.
        text_preview = ""
        try:
            f.stream.seek(0)
        except Exception:
            pass
        saved_doc = pipeline.save_uploaded_file(user["id"], f)
        text_preview = extract.extract_text(saved_doc["stored_path"])

        if not pipeline.matches_confirmed_owner(user["id"], saved_doc["filename"], text_preview):
            db.delete_document(user["id"], saved_doc["id"])
            identity_rejected.append(saved_doc["filename"])
            continue

        saved.append(saved_doc)

    if identity_rejected and not saved:
        return jsonify(
            {
                "ok": False,
                "error": (
                    "These documents don't appear to belong to "
                    f"{user.get('confirmed_owner_name') or 'the account holder'}: "
                    f"{', '.join(identity_rejected)}. Combining different people's documents "
                    "in one account isn't allowed."
                ),
            }
        ), 400

    if not saved:
        return jsonify({"ok": False, "error": "None of the selected files are a supported type."}), 400

    try:
        pipeline.run_analysis_for_user(user["id"])
    except analyzer.CVAnalyzerError as e:
        # The files are saved either way — only the AI scoring failed
        # (e.g. missing API key) — so tell the user the upload worked
        # but scoring couldn't run, rather than pretending nothing
        # happened.
        return jsonify(
            {
                "ok": True,
                "warning": f"Files uploaded, but scoring couldn't run: {e}",
                "rejected": rejected,
                "state": pipeline.get_dashboard_state(user["id"]),
            }
        )

    # Store score delta + per-dimension deltas on each newly saved document
    _post_row = db.get_latest_analysis(user["id"])
    _post_analysis = json.loads(_post_row["result_json"]) if _post_row else {}
    _score_after = float(_post_analysis.get("overall_score") or 0)
    _delta = round(_score_after - _score_before, 2)
    _dims_after = {d["label"]: d["score"] for d in (_post_analysis.get("dimensions") or [])}
    _dim_deltas = {
        lbl: round(_dims_after.get(lbl, 0) - _dims_before.get(lbl, 0), 2)
        for lbl in set(list(_dims_before) + list(_dims_after))
        if round(_dims_after.get(lbl, 0) - _dims_before.get(lbl, 0), 2) != 0
    }
    if _delta != 0 or _dim_deltas:
        _dim_json = json.dumps(_dim_deltas) if _dim_deltas else None
        for _doc in saved:
            db.set_document_score_delta(user["id"], _doc["id"], _delta, _dim_json)

    warning = None
    if identity_rejected:
        warning = f"Skipped (doesn't match this account's confirmed identity): {', '.join(identity_rejected)}"

    response = {"ok": True, "rejected": rejected, "state": pipeline.get_dashboard_state(user["id"])}
    if warning:
        response["warning"] = warning
    return jsonify(response)


@app.route("/api/documents/<int:document_id>", methods=["DELETE"])
@auth.login_required
def api_delete_document(document_id):
    user = auth.current_user()
    db.delete_document(user["id"], document_id)
    remaining = db.get_documents_for_user(user["id"])
    if not remaining:
        db.reset_documents_confirmed(user["id"])
        return jsonify({"ok": True, "state": pipeline.get_dashboard_state(user["id"])})
    try:
        pipeline.run_analysis_for_user(user["id"])
    except analyzer.CVAnalyzerError as e:
        return jsonify({
            "ok": True,
            "warning": f"Document removed, but scoring couldn't re-run: {e}",
            "state": pipeline.get_dashboard_state(user["id"]),
        })
    return jsonify({"ok": True, "state": pipeline.get_dashboard_state(user["id"])})


@app.route("/api/reanalyze", methods=["POST"])
@auth.login_required
def api_reanalyze():
    user = auth.current_user()
    extra_context = request.json.get("extra_context", "") if request.is_json else ""
    try:
        pipeline.run_analysis_for_user(user["id"], extra_context=extra_context)
    except analyzer.CVAnalyzerError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    return jsonify({"ok": True, "state": pipeline.get_dashboard_state(user["id"])})


@app.route("/api/skills", methods=["POST"])
@auth.login_required
def api_add_skill():
    user = auth.current_user()
    data = request.json if request.is_json else request.form
    label = (data.get("label") or "").strip()
    if not label:
        return jsonify({"ok": False, "error": "Skill cannot be empty."}), 400
    if len(label) > 60:
        return jsonify({"ok": False, "error": "Skill is too long."}), 400
    db.add_skill(user["id"], label, source="manual")
    return jsonify({"ok": True, "state": pipeline.get_dashboard_state(user["id"])})


@app.route("/api/skills/<int:skill_id>", methods=["DELETE"])
@auth.login_required
def api_delete_skill(skill_id):
    user = auth.current_user()
    db.delete_skill(user["id"], skill_id)
    return jsonify({"ok": True, "state": pipeline.get_dashboard_state(user["id"])})


@app.route("/api/profile", methods=["POST"])
@auth.login_required
def api_update_profile():
    user = auth.current_user()
    data = request.json if request.is_json else request.form
    fields = {}
    for key in ("full_name", "headline", "email", "location", "phone"):
        if key in data:
            fields[key] = (data.get(key) or "").strip()

    if "email" in fields:
        try:
            fields["email"] = auth.validate_email(fields["email"])
        except auth.AuthError as e:
            return jsonify({"ok": False, "error": str(e)}), 400
        existing = db.get_user_by_email(fields["email"])
        if existing and existing["id"] != user["id"]:
            return jsonify({"ok": False, "error": "That email is already in use by another account."}), 400
    if "full_name" in fields:
        try:
            fields["full_name"] = auth.validate_full_name(fields["full_name"])
        except auth.AuthError as e:
            return jsonify({"ok": False, "error": str(e)}), 400

    db.update_profile_fields(user["id"], **fields)
    return jsonify({"ok": True, "state": pipeline.get_dashboard_state(user["id"])})


@app.route("/api/profile/password", methods=["POST"])
@auth.login_required
def api_change_password():
    user = auth.current_user()
    data = request.get_json(force=True)
    try:
        new_security_key = auth.change_password(
            user["id"],
            data.get("current_password_or_key") or "",
            data.get("new_password") or "",
            data.get("confirm_password") or "",
        )
        # Only non-None when the security key (rather than the ordinary
        # password) was used as proof -- it's rotated on use, so the
        # caller needs the new one to show the user once.
        return jsonify({"ok": True, "new_security_key": new_security_key})
    except auth.AuthError as e:
        return jsonify({"ok": False, "error": str(e)}), 400


@app.route("/api/profile/security-key/regenerate", methods=["POST"])
@auth.login_required
def api_regenerate_security_key():
    user = auth.current_user()
    new_security_key = auth.regenerate_security_key(user["id"])
    return jsonify({"ok": True, "security_key": new_security_key})


@app.route("/api/profile/documents")
@auth.login_required
def api_profile_documents():
    user = auth.current_user()
    docs = _real_uploaded_docs(user["id"])
    return jsonify({
        "ok": True,
        "documents": [
            {
                "id": d["id"],
                "filename": d["filename"],
                "category": d.get("category", ""),
                "file_type": d.get("file_type", ""),
                "file_size": d.get("file_size"),
                "uploaded_at": d.get("uploaded_at"),
            }
            for d in docs
        ],
    })


@app.route("/profile")
@auth.login_required
def profile_page():
    user = auth.current_user()
    return render_template(
        "profile.html", user=user,
        avatar_url=_avatar_url_for(user), initials=_initials(user.get("full_name") or ""),
    )


@app.route("/api/friends/invite", methods=["POST"])
@auth.login_required
def api_friends_invite():
    user = auth.current_user()
    data = request.json if request.is_json else request.form
    username = (data.get("username") or "").strip()
    if not username:
        return jsonify({"ok": False, "error": "Enter a username."}), 400

    target = db.get_user_by_username(username)
    if not target:
        return jsonify({"ok": False, "error": "No account found with that username."}), 404
    if target["id"] == user["id"]:
        return jsonify({"ok": False, "error": "You can't add yourself."}), 400

    existing = db.get_active_friend_request_between(user["id"], target["id"])
    if existing:
        if existing["status"] == "accepted":
            return jsonify({"ok": False, "error": "You're already friends."}), 400
        return jsonify({"ok": False, "error": "A friend request is already pending."}), 400

    db.create_friend_request(user["id"], target["id"])
    return jsonify({"ok": True})


@app.route("/api/friends/requests")
@auth.login_required
def api_friends_requests():
    user = auth.current_user()
    return jsonify({"requests": db.get_pending_incoming_requests(user["id"])})


@app.route("/api/friends/requests/<int:request_id>/respond", methods=["POST"])
@auth.login_required
def api_friends_respond(request_id):
    user = auth.current_user()
    data = request.json if request.is_json else request.form
    accept = bool(data.get("accept"))

    req = db.get_friend_request_by_id(request_id)
    if not req or req["to_user_id"] != user["id"] or req["status"] != "pending":
        return jsonify({"ok": False, "error": "Request not found."}), 404

    db.respond_to_friend_request(request_id, "accepted" if accept else "declined")
    return jsonify({"ok": True})


@app.route("/api/friends")
@auth.login_required
def api_friends_list():
    user = auth.current_user()
    return jsonify({"friends": db.get_friends_for_user(user["id"])})


# Avatars render at most ~64px on screen (sidebar/profile), so a phone
# photo straight off the camera (often several MB) would otherwise get
# base64-embedded at full resolution into every page that shows it --
# dashboard, profile, every sidebar render. Downscaling once at upload
# time keeps that inline payload small permanently.
AVATAR_MAX_DIMENSION = 256


@app.route("/api/profile/photo", methods=["POST"])
@auth.login_required
def api_upload_avatar():
    import base64
    import io
    from PIL import Image, ImageOps

    user = auth.current_user()
    f = request.files.get("photo")
    if not f or f.filename == "":
        return jsonify({"ok": False, "error": "No file selected."}), 400
    ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        return jsonify({"ok": False, "error": "Unsupported image type."}), 400

    try:
        image = Image.open(f.stream)
        image = ImageOps.exif_transpose(image)  # respect the camera's rotation before resizing
        image.thumbnail((AVATAR_MAX_DIMENSION, AVATAR_MAX_DIMENSION), Image.LANCZOS)
        if image.mode != "RGB":
            image = image.convert("RGB")
        buf = io.BytesIO()
        image.save(buf, format="JPEG", quality=82)
        encoded = base64.b64encode(buf.getvalue()).decode("ascii")
        data_uri = f"data:image/jpeg;base64,{encoded}"
    except Exception:
        return jsonify({"ok": False, "error": "Couldn't read that image."}), 400

    db.update_profile_fields(user["id"], avatar_path=data_uri)
    # No "state" field here -- the caller (profile.js) only ever reads
    # avatar_url, so computing the full multi-query dashboard state on
    # every photo upload was pure wasted work.
    return jsonify({"ok": True, "avatar_url": data_uri})


@app.route("/api/account/delete", methods=["POST"])
@auth.login_required
def api_delete_account():
    """
    Permanently deletes the logged-in user's account. The DB row delete
    cascades to every other table (documents, skills, analyses, etc. —
    see db.delete_user's docstring), but that cascade doesn't touch the
    filesystem, so the user's uploaded-files folder is removed here too.
    """
    import shutil
    user = auth.current_user()
    user_id = user["id"]
    upload_dir = pipeline.user_upload_dir(user_id)
    db.delete_user(user_id)
    shutil.rmtree(upload_dir, ignore_errors=True)
    auth.log_out_user()
    return jsonify({"ok": True})


@app.route("/api/clear-cache", methods=["POST"])
@auth.login_required
def api_clear_cache():
    user = auth.current_user()
    cache_module.clear_cache()
    db.clear_analyses(user["id"])
    return jsonify({"ok": True, "state": pipeline.get_dashboard_state(user["id"])})


@app.route("/api/applications", methods=["POST"])
@auth.login_required
def api_add_application():
    user = auth.current_user()
    data = request.json if request.is_json else request.form
    job_title = (data.get("job_title") or "").strip()
    company = (data.get("company") or "").strip()
    if not job_title or not company:
        return jsonify({"ok": False, "error": "Job title and company are required."}), 400
    db.add_application(user["id"], job_title, company)
    return jsonify({"ok": True, "state": pipeline.get_dashboard_state(user["id"])})


@app.route("/privacy-policy")
def privacy_policy():
    return render_template("privacy_policy.html")


@app.route("/terms")
@auth.login_required
def terms_page():
    user = auth.current_user()
    return render_template("terms_gate.html", already_accepted=bool(user.get("disclaimer_accepted")))


@app.route("/api/disclaimer-accept", methods=["POST"])
@auth.login_required
def api_disclaimer_accept():
    user = auth.current_user()
    db.set_disclaimer_accepted(user["id"])
    return jsonify({"ok": True})


@app.route("/api/target-field", methods=["POST"])
@auth.login_required
def api_set_target_field():
    user = auth.current_user()
    data = request.json if request.is_json else request.form
    target_field = (data.get("target_field") or "").strip()
    db.set_target_field(user["id"], target_field)
    return jsonify({"ok": True, "state": pipeline.get_dashboard_state(user["id"])})


@app.route("/api/custom-instructions", methods=["POST"])
@auth.login_required
def api_set_custom_instructions():
    user = auth.current_user()
    data = request.get_json(force=True)
    custom_instructions = (data.get("custom_instructions") or "").strip()[:3000]
    db.set_custom_instructions(user["id"], custom_instructions)
    return jsonify({"ok": True})


@app.route("/api/settings/remember-chats", methods=["POST"])
@auth.login_required
def api_set_remember_chats():
    user = auth.current_user()
    data = request.get_json(force=True)
    db.set_remember_all_chats(user["id"], bool(data.get("enabled")))
    return jsonify({"ok": True})


AI_UPLOAD_ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "instance", "ai_uploads")


@app.route("/api/chat/upload", methods=["POST"])
@auth.login_required
def api_chat_upload():
    """Upload a file for use in AI chat. Stored server-side, never added to CV profile."""
    user = auth.current_user()
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "No file provided"}), 400
    f = request.files["file"]
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "Empty file"}), 400

    user_upload_dir = os.path.join(AI_UPLOAD_ROOT, str(user["id"]))
    os.makedirs(user_upload_dir, exist_ok=True)

    safe_name = uuid.uuid4().hex + "_" + f.filename.replace(" ", "_")
    stored_path = os.path.join(user_upload_dir, safe_name)
    f.save(stored_path)

    mime_type = f.content_type or ""
    text_content = ""
    if not mime_type.startswith("image/"):
        try:
            text_content = extract.extract_text(stored_path)
            if text_content:
                text_content = text_content.strip()[:8000]
        except Exception:
            pass

    att_id = db.add_chat_attachment(user["id"], f.filename, stored_path, mime_type, text_content)
    return jsonify({"ok": True, "id": att_id, "name": f.filename, "mime": mime_type, "is_image": mime_type.startswith("image/")})


@app.route("/api/chat/conversations", methods=["GET"])
@auth.login_required
def api_get_conversations():
    user = auth.current_user()
    convs = db.get_conversations_for_user(user["id"])
    return jsonify({"ok": True, "conversations": convs})


def _generate_conversation_title(messages):
    """First-exchange AI naming for a brand new chat thread ("auto-named
    by AI from first message"). Returns None on any failure (including
    no API key configured) so the caller can fall back to its own
    truncated-text title instead of leaving a new chat untitled."""
    from openai import OpenAI
    convo_snippet = "\n".join(
        f"{m.get('role')}: {(m.get('text') or '')[:400]}" for m in messages[:4]
    )
    if not convo_snippet.strip():
        return None
    try:
        # A short, no-retry timeout on purpose: this is a best-effort
        # cosmetic nicety with an already-good fallback (the caller's own
        # truncated-text title), not a core feature worth making the user
        # wait on. The generous 90s/3-retry budget used for real scoring
        # work would otherwise leave a slow or unreachable API holding up
        # the very first save of a brand new conversation for minutes.
        client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=8.0, max_retries=0)
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Give this chat a short title, 3-6 words, plain text, no quotes or trailing punctuation, summarizing what it's about."},
                {"role": "user", "content": convo_snippet},
            ],
            max_tokens=20,
        )
        title = (resp.choices[0].message.content or "").strip().strip('"').strip("'")
        return title[:80] if title else None
    except Exception:
        return None


@app.route("/api/chat/conversations", methods=["POST"])
@auth.login_required
def api_save_conversation():
    """Create a new conversation or append messages to an existing one.
    The title is set once, at creation (AI-named from the first
    exchange when possible) -- later saves never touch it, so it can't
    drift back to raw truncated text on a later message, and promoted
    job threads keep the "{Role} · {Company}" title promotion set."""
    user = auth.current_user()
    data = request.get_json(force=True)
    conv_id = data.get("conversation_id")
    fallback_title = (data.get("title") or "Conversation")[:80]
    messages = data.get("messages", [])

    if not conv_id:
        title = _generate_conversation_title(messages) or fallback_title
        conv_id = db.create_conversation(user["id"], title)
    elif not db.conversation_belongs_to_user(conv_id, user["id"]):
        # conversation_id came from the client -- without this check, a
        # crafted id belonging to another user would let their entire
        # conversation be wiped and overwritten by the DELETE+INSERT below.
        return jsonify({"ok": False, "error": "Conversation not found."}), 404

    # Replace all messages for this conversation (simplest sync strategy)
    conn = db.get_db()
    try:
        conn.execute("DELETE FROM chat_messages WHERE conversation_id = ?", (conv_id,))
        conn.commit()
        for m in messages:
            db.add_chat_message(conv_id, m.get("role", "user"), m.get("text", ""), m.get("attachment_ids", []), m.get("card"))
        db.touch_conversation(conv_id, user["id"])
    finally:
        conn.close()

    return jsonify({"ok": True, "conversation_id": conv_id})


def _serialize_conversation_messages(conv_id, user_id):
    msgs = db.get_messages_for_conversation(conv_id, user_id)
    result = []
    for m in msgs:
        att_ids = json.loads(m.get("attachment_ids_json") or "[]")
        card = json.loads(m["card_json"]) if m.get("card_json") else None
        result.append({"role": m["role"], "text": m["text"], "attachment_ids": att_ids, "card": card})
    return result


@app.route("/api/chat/conversations/<int:conv_id>", methods=["GET"])
@auth.login_required
def api_get_conversation(conv_id):
    user = auth.current_user()
    return jsonify({"ok": True, "messages": _serialize_conversation_messages(conv_id, user["id"])})


@app.route("/api/chat/conversations/<int:conv_id>", methods=["DELETE"])
@auth.login_required
def api_delete_conversation(conv_id):
    user = auth.current_user()
    db.delete_conversation(conv_id, user["id"])
    return jsonify({"ok": True})


@app.route("/api/chat/conversations", methods=["DELETE"])
@auth.login_required
def api_delete_all_conversations():
    user = auth.current_user()
    db.delete_all_conversations_for_user(user["id"])
    return jsonify({"ok": True})


@app.route("/api/chat/conversations/<int:conv_id>/promote", methods=["POST"])
@auth.login_required
def api_promote_conversation(conv_id):
    """The job-thread promotion mechanic: called the moment a Verdict or
    Document card lands in a chat. Title snaps to "{Role} · {Company}",
    the thread gets a live status badge, and it sorts above plain chats
    -- this IS the tracker, there's no separate screen for it."""
    user = auth.current_user()
    data = request.get_json(force=True)
    fit_score = data.get("fit_score")
    try:
        fit_score = int(fit_score) if fit_score is not None else None
    except (TypeError, ValueError):
        fit_score = None
    db.promote_conversation(
        conv_id, user["id"],
        job_title=(data.get("job_title") or "").strip()[:200],
        company=(data.get("company") or "").strip()[:200],
        fit_score=fit_score,
        status_label=(data.get("status_label") or "").strip()[:100],
    )
    return jsonify({"ok": True})


@app.route("/api/chat/conversations/<int:conv_id>/status", methods=["POST"])
@auth.login_required
def api_update_conversation_status(conv_id):
    """Updates a job thread's live status badge -- e.g. the Document
    card's "Mark as applied" action moving a thread from "CV ready" to
    "Sent"."""
    user = auth.current_user()
    data = request.get_json(force=True)
    db.update_conversation_status(conv_id, user["id"], (data.get("status_label") or "").strip()[:100])
    return jsonify({"ok": True})


@app.route("/api/chat/attachment-thumb/<int:att_id>")
@auth.login_required
def api_chat_attachment_thumb(att_id):
    """Serve the stored image file for preview thumbnails."""
    from flask import send_file
    user = auth.current_user()
    att = db.get_chat_attachment(user["id"], att_id)
    if not att or not att["mime_type"].startswith("image/"):
        return ("Not found", 404)
    return send_file(att["stored_path"], mimetype=att["mime_type"])


@app.route("/api/cv-text")
@auth.login_required
def api_cv_text():
    """Return the user's CV as HTML for the editor, extracting text live from the stored file."""
    import extract as _extract
    user = auth.current_user()
    documents = db.get_documents_for_user(user["id"])
    if not documents:
        return jsonify({"html": None})
    # Prefer PDFs/docs; fall back to all documents
    cv_docs = [d for d in documents if d.get("file_type", "").lower() in ("pdf", "docx", "doc", "cv", "resume")]
    ordered = cv_docs or documents
    combined = ""
    for d in ordered:
        text = (d.get("text_content") or "").strip()
        if not text:
            try:
                text = _extract.extract_text(d["stored_path"]).strip()
            except Exception:
                pass
        if text:
            combined = text
            break  # use the first successful CV
    if not combined:
        return jsonify({"html": None})
    # Convert to HTML: detect headings (short ALL-CAPS lines) and paragraphs
    html_parts = []
    for line in combined.split("\n"):
        stripped = line.strip()
        if not stripped:
            html_parts.append("<p><br></p>")
        elif len(stripped) < 60 and stripped.isupper():
            html_parts.append(f"<p><strong>{stripped}</strong></p>")
        else:
            html_parts.append(f"<p>{stripped}</p>")
    return jsonify({"html": "".join(html_parts)})


@app.route("/api/document-insight/<int:doc_id>")
@auth.login_required
def api_document_insight(doc_id):
    """AI explanation of exactly what a specific document contributed to the profile."""
    from openai import OpenAI
    import extract as _extract
    user = auth.current_user()

    docs = db.get_documents_for_user(user["id"])
    doc = next((d for d in docs if d["id"] == doc_id), None)
    if not doc:
        return jsonify({"error": "Document not found."}), 404

    # Return cached insight if available (only if it's HTML — plain text cache is stale)
    cached = doc.get("insight_cache") or ""
    if cached and ("<p>" in cached or "<ul>" in cached or "<strong>" in cached):
        return jsonify({"insight": cached, "filename": doc["filename"],
                        "score_delta": doc.get("score_delta"),
                        "dim_deltas": json.loads(doc["dimension_deltas"]) if doc.get("dimension_deltas") else {}})

    # Extract text
    text = (doc.get("text_content") or "").strip()
    if not text:
        try:
            text = _extract.extract_text(doc["stored_path"]).strip()
        except Exception:
            text = ""

    _analysis_row = db.get_latest_analysis(user["id"])
    _analysis_data = json.loads(_analysis_row["result_json"]) if _analysis_row and _analysis_row.get("result_json") else {}
    skills = db.get_skills_for_user(user["id"])
    skill_names = [s["label"] if isinstance(s, dict) else str(s) for s in skills]
    dimensions = _analysis_data.get("dimensions") or []
    overall = float(_analysis_data.get("overall_score") or 0)
    delta = float(doc.get("score_delta") or 0)
    dim_deltas = json.loads(doc["dimension_deltas"]) if doc.get("dimension_deltas") else {}
    all_filenames = [d["filename"] for d in docs if d["id"] != doc_id]

    client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
    system = (
        "You are a career intelligence analyst inside the Employable platform. "
        "Respond ONLY with valid HTML — use <p>, <ul>, <li>, <strong> tags. "
        "Never use asterisks (*) or markdown. Keep paragraphs short (2-3 sentences max). "
        "Use bullet lists for multiple items. Be specific and reference real content from the document.\n\n"
        "Structure your response as:\n"
        "1. A short intro paragraph — what this document proves to an employer\n"
        "2. A <ul> list of which score dimensions it raised and by how much (use the dim_deltas provided)\n"
        "3. A short paragraph on how it complements the other documents\n"
        "4. One concrete <strong>improvement tip</strong> to make this document stronger\n\n"
        "Be truthful, specific, and genuinely helpful. Never be vague."
    )
    dim_lines = "\n".join(f"  {d['label']}: {d['score']}/10" for d in dimensions)
    dim_delta_lines = "\n".join(f"  {lbl}: +{v} pts" for lbl, v in dim_deltas.items()) if dim_deltas else "  (no specific dimension data)"
    other_docs = ", ".join(all_filenames) if all_filenames else "none"
    prompt = (
        f"Document: {doc['filename']} ({doc['file_type'].upper()})\n"
        f"Overall score impact: +{delta} pts (was {round(overall - delta, 1)} → now {overall})\n\n"
        f"Dimension-level score increases from this document:\n{dim_delta_lines}\n\n"
        f"Current dimension scores:\n{dim_lines}\n\n"
        f"Document content:\n{text[:3000]}\n\n"
        f"Other documents in profile: {other_docs}\n"
        f"Registered skills: {', '.join(skill_names[:20])}\n\n"
        f"Explain specifically what '{doc['filename']}' contributes. Reference the dim_deltas in your bullet list."
    )
    try:
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            max_tokens=700,
        )
        insight = resp.choices[0].message.content.strip()
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    db.set_document_insight_cache(user["id"], doc_id, insight)
    return jsonify({"insight": insight, "filename": doc["filename"], "score_delta": delta, "dim_deltas": dim_deltas})


@app.route("/api/cv-text-upload", methods=["POST"])
@auth.login_required
def api_cv_text_upload():
    """Accept a CV file uploaded directly from the user's device, return as editor HTML."""
    import extract as _extract, tempfile, os as _os
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify({"error": "No file received."}), 400
    ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else "txt"
    if ext not in ("pdf", "doc", "docx", "txt", "rtf"):
        return jsonify({"error": "Unsupported file type. Please upload a PDF, DOCX, or TXT."}), 400
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            f.save(tmp.name)
            tmp_path = tmp.name
        text = _extract.extract_text(tmp_path).strip()
    except Exception as e:
        return jsonify({"error": f"Could not read file: {e}"}), 500
    finally:
        try: _os.unlink(tmp_path)
        except Exception: pass
    if not text:
        return jsonify({"error": "No readable text found in that file."}), 400
    html_parts = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            html_parts.append("<p><br></p>")
        elif len(stripped) < 60 and stripped.isupper():
            html_parts.append(f"<p><strong>{stripped}</strong></p>")
        else:
            html_parts.append(f"<p>{stripped}</p>")
    return jsonify({"html": "".join(html_parts)})


@app.route("/api/roadmap/completions")
@auth.login_required
def api_roadmap_completions():
    user = auth.current_user()
    completions = db.get_roadmap_completions(user["id"])
    return jsonify({"completions": [dict(c) for c in completions]})


@app.route("/api/roadmap/complete", methods=["POST"])
@auth.login_required
def api_roadmap_complete():
    """
    AI evaluates whether a selected document fulfills a roadmap
    objective. The AI only judges fulfilled/not-fulfilled (plus a
    reason and, if not fulfilled, concrete next steps) — it does NOT
    invent its own point value. The point value awarded is whatever
    the dashboard already displayed for this item (computed in
    pipeline._filter_and_recompute_roadmap using the exact same formula
    as the visible gauge/bars), passed through from the frontend and
    clamped to a sane range here. This is what guarantees the number
    the user saw on the card is the exact number they get — no second,
    independently-invented number from a different AI call.
    """
    from openai import OpenAI
    import extract as _extract
    user = auth.current_user()
    data = request.get_json(force=True)
    item_label = (data.get("item_label") or "").strip()
    item_description = (data.get("item_description") or "").strip()
    doc_id = data.get("doc_id")
    try:
        # 2.0 is the maximum headroom any single visible dimension can
        # have (see DASHBOARD_VISIBLE_DIMENSIONS scoring) — anything
        # outside that range could only be a bad/tampered request.
        displayed_points = max(0.0, min(2.0, float(data.get("points") or 0)))
    except (TypeError, ValueError):
        displayed_points = 0.0
    if not item_label or not doc_id:
        return jsonify({"error": "Missing item_label or doc_id"}), 400

    # Check not already completed
    completions = db.get_roadmap_completions(user["id"])
    if any(c["item_label"] == item_label for c in completions):
        return jsonify({"error": "already_completed"}), 400

    docs = db.get_documents_for_user(user["id"])
    doc = next((d for d in docs if d["id"] == int(doc_id)), None)
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    text = (doc.get("text_content") or "").strip()
    if not text:
        try:
            text = _extract.extract_text(doc["stored_path"]).strip()
        except Exception:
            text = ""

    _analysis_row = db.get_latest_analysis(user["id"])
    _analysis_data = json.loads(_analysis_row["result_json"]) if _analysis_row and _analysis_row.get("result_json") else {}
    overall = float(_analysis_data.get("overall_score") or 0)

    client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
    system = (
        "You are a strict but fair career coach evaluating whether a document fulfills a specific career improvement objective. "
        "Reply ONLY with a JSON object in this exact format:\n"
        '{"fulfilled": true/false, "reason": "one short sentence explaining your verdict", "steps": []}\n'
        "Rules:\n"
        "- steps: if NOT fulfilled, provide 3–5 specific, actionable steps (as an array of strings) that tell the user exactly "
        "what they need to do or add to their document to satisfy this objective. Be specific to the objective — not generic advice. "
        "If fulfilled, steps should be an empty array [].\n"
        "Be truthful — if the document genuinely demonstrates the required skill or output, mark it fulfilled."
    )
    prompt = (
        f"Roadmap objective: {item_label}\n"
        f"Objective detail: {item_description}\n\n"
        f"Document submitted: {doc['filename']}\n"
        f"Document content (first 2500 chars):\n{text[:2500]}\n\n"
        f"Evaluate whether this document fulfills the roadmap objective. Reply with JSON only."
    )
    try:
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            max_tokens=200,
            temperature=0.2,
            response_format={"type": "json_object"},
        )
        result = json.loads(resp.choices[0].message.content)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    fulfilled = bool(result.get("fulfilled"))
    reason = result.get("reason", "")
    steps = result.get("steps") or []
    points = displayed_points if fulfilled else 0.0

    if fulfilled and points > 0:
        db.add_roadmap_completion(user["id"], item_label, int(doc_id), points)

    return jsonify({
        "fulfilled": fulfilled,
        "reason": reason,
        "points": points,
        "steps": steps,
        "overall_before": overall,
    })


@app.route("/api/cv-edit", methods=["POST"])
@auth.login_required
def api_cv_edit():
    """AI edits the CV content based on a user instruction. Works with HTML for precise formatting."""
    from openai import OpenAI
    user = auth.current_user()
    data = request.get_json(force=True)
    instruction = (data.get("instruction") or "").strip()
    cv_html = (data.get("cv_html") or "").strip()
    cv_content = (data.get("cv_content") or "").strip()
    if not instruction:
        return jsonify({"error": "Missing instruction."}), 400

    # Same grounding the main chat endpoint uses (see api_chat) -- without
    # this, an empty editor + an instruction like "write me a CV like
    # mine" had nothing real to draw from, and the model would invent a
    # complete fake person (wrong job history, wrong employers, wrong
    # everything) rather than ever admitting it had nothing to go on.
    state = pipeline.get_dashboard_state(user["id"])
    profile = state.get("profile", {})
    full_docs = db.get_documents_for_user(user["id"])
    doc_texts = []
    total_chars = 0
    DOC_CHAR_CAP = 3000
    TOTAL_CHAR_CAP = 20000
    for d in full_docs:
        if total_chars >= TOTAL_CHAR_CAP:
            break
        try:
            txt = (d.get("content") or "").strip()
            if not txt and d.get("stored_path") and os.path.exists(d["stored_path"]):
                txt = (extract.extract_text(d["stored_path"]) or "").strip()
            if txt:
                snippet = txt[:DOC_CHAR_CAP]
                doc_texts.append(f"[{d['filename']}]\n{snippet}")
                total_chars += len(snippet)
        except Exception:
            pass
    doc_content_block = "\n\n---\n\n".join(doc_texts) if doc_texts else "No documents uploaded yet."
    skill_names = [s["label"] for s in state.get("skills", [])]

    import re as _re
    system = (
        "You are an expert CV/document editor built into the Employable platform. "
        "CRITICAL: Always attempt to understand and fulfill the user's intent, even if their instruction contains spelling mistakes, typos, or imprecise phrasing. "
        "Never refuse, do nothing, or ask for clarification — silently make your best reasonable interpretation and act on it. "
        "OUTPUT RULES — you MUST follow these exactly:\n"
        "- Return a JSON object with exactly two keys: 'html' and 'description'.\n"
        "- 'html': the full updated document as valid HTML. Use <p>, <strong>, <em>, <ul>, <li>, <h1>, <h2>, "
        "<h3>, <hr>, <br> tags, plus the exact entry-row pattern described below for job/education entries. "
        "NEVER use markdown asterisks, hyphens for bullets, --- separators, or backticks. "
        "No <html>, <head>, <body> wrappers. No code fences.\n"
        "- 'description': one short, specific sentence (max 20 words) describing exactly what you changed — e.g. "
        "'Made all section headings bold and centered the name at the top.' "
        "Be specific to what was actually done, not generic.\n\n"
        "DESIGN — when writing or restructuring a CV (not just tweaking a sentence), apply real CV design "
        "sense, not just correct facts in a wall of text: a clear name/contact header, short bold section "
        "headings (Experience, Education, Skills, etc.) in a consistent order, reverse-chronological entries "
        "within each section, bullet points for achievements rather than dense paragraphs, consistent bold "
        "usage for role titles or employer names (pick one convention and stick to it), and generous <hr>/"
        "spacing between sections so it's scannable in a 6-second recruiter skim. The app applies its own "
        "visual template on top of your HTML, so don't hand-roll layout styling yourself -- just follow this "
        "structure:\n"
        f"{CV_STRUCTURE_GUIDANCE}\n\n"
        "GROUNDING — this is the most important rule, above all others: every fact in the document "
        "(employer names, job titles, dates, schools, achievements, contact details) must come from "
        "the user's real profile/documents below, or already be present in the current document HTML. "
        "NEVER invent a person, career, employer, achievement, or biography that isn't actually theirs "
        "-- not even a plausible-sounding placeholder one. If the document is empty and the user asks you "
        "to write or generate a CV, build it FROM the real document content below. If there isn't enough "
        "real information to do that (no documents uploaded, or nothing usable in them), say so plainly "
        "in 'description' and return the document unchanged rather than fabricating a fake CV.\n\n"
        f"This user's real profile:\n"
        f"- Name: {profile.get('full_name') or 'Unknown'}\n"
        f"- Location: {profile.get('location') or 'Not specified'}\n"
        f"- Skills: {', '.join(skill_names) if skill_names else 'None listed.'}\n\n"
        f"Full content of this user's actual uploaded documents (the ONLY source of truth for any CV "
        f"content you write):\n{doc_content_block}"
    )
    prompt = f"Current document HTML:\n{cv_html if cv_html else '(empty — generate fresh content)'}\n\nInstruction: {instruction}"

    try:
        client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            max_tokens=3200,
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content.strip()
        try:
            parsed = json.loads(raw)
            updated = parsed.get("html", "")
            description = parsed.get("description", "Done.")
        except Exception:
            updated = raw
            description = "Done."
        # Strip accidental code fences
        if updated.startswith("```"):
            updated = updated.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        # Convert any markdown that slipped through into HTML
        updated = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', updated)
        updated = _re.sub(r'\*(.+?)\*', r'<em>\1</em>', updated)
        updated = _re.sub(r'(?m)^---+\s*$', '<hr>', updated)
        return jsonify({"updated_html": updated, "description": description})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/cv-download/docx", methods=["POST"])
@auth.login_required
def api_cv_download_docx():
    """Generate a DOCX file that reflects the editor's real formatting."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from flask import send_file
    import io
    from cv_export import parse_cv_html

    data = request.get_json(force=True)
    cv_html = (data.get("cv_html") or "").strip()
    content = (data.get("content") or "").strip()
    blocks = parse_cv_html(cv_html) if cv_html else []
    if not blocks and not content:
        return jsonify({"error": "No content provided."}), 400

    ALIGN_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    HEADING_SIZE = {"h1": 16, "h2": 13, "h3": 12}
    DOCX_MARGINS = {
        "narrow": (0.6, 0.6),
        "normal": (1.0, 1.15),
        "wide": (1.4, 1.6),
    }
    v_in, h_in = DOCX_MARGINS.get(data.get("margins"), DOCX_MARGINS["normal"])

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(v_in)
        section.bottom_margin = Inches(v_in)
        section.left_margin = Inches(h_in)
        section.right_margin = Inches(h_in)

    from docx.enum.text import WD_TAB_ALIGNMENT

    # Stacked entry-line paragraphs (new AI generations) -- same styling
    # intent as the PDF's ENTRY_LINE_STYLE, expressed as docx run attrs.
    ENTRY_LINE_STYLE = {
        "cv-entry-title": {"size": 11.5, "bold": True, "space_before": 9, "color": None},
        "cv-entry-company": {"size": 11, "bold": True, "space_before": 0, "color": None},
        "cv-entry-dates": {"size": 10, "bold": True, "space_before": 0, "color": "444444"},
        "cv-entry-reference": {"size": 9.5, "bold": False, "space_before": 1, "color": "777777"},
    }

    if blocks:
        for block in blocks:
            if block["kind"] == "hr":
                p = doc.add_paragraph("─" * 40)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            if block["kind"] == "header":
                # Name on the left, job title + stacked contact lines
                # right-aligned -- a right tab stop plus explicit line
                # breaks (add_run can't embed "\n", Word needs a real
                # break run) gets the same two-column shape as the PDF's
                # header table without needing an actual docx table.
                usable_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                p = doc.add_paragraph()
                p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
                name_run = p.add_run(block["name_text"])
                name_run.font.name = "Arial"
                name_run.font.size = Pt(20)
                name_run.bold = True
                right_lines = []
                if block["title_text"]:
                    right_lines.append((block["title_text"], True, 13, None))
                for contact_line in block["contact_text"].split("\n"):
                    if contact_line:
                        right_lines.append((contact_line, False, 10.5, "666666"))
                if right_lines:
                    p.add_run("\t")
                    for i, (text, bold, size, color) in enumerate(right_lines):
                        if i:
                            p.add_run().add_break()
                        run = p.add_run(text)
                        run.font.name = "Arial"
                        run.font.size = Pt(size)
                        run.bold = bold
                        if color:
                            run.font.color.rgb = RGBColor.from_string(color)
                continue

            if block["kind"] == "row":
                # Same title-left/date-right entry line as the PDF, done
                # here with a right tab stop instead of a table -- python-docx
                # can't easily add a borderless 2-cell table without extra
                # ceremony, and a tab stop is the standard Word technique
                # for this exact "left text ... right text" pattern anyway.
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(8)
                usable_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
                title_run = p.add_run(block["title_text"])
                title_run.font.name = "Arial"
                title_run.font.size = Pt(11)
                title_run.bold = True
                if block["date_text"]:
                    p.add_run("\t")
                    date_run = p.add_run(block["date_text"])
                    date_run.font.name = "Arial"
                    date_run.font.size = Pt(10)
                    date_run.italic = True
                continue

            p = doc.add_paragraph(style="List Bullet" if (block["kind"] == "li" and not block["ordered"]) else
                                         "List Number" if (block["kind"] == "li" and block["ordered"]) else None)
            p.alignment = ALIGN_MAP.get(block["align"], WD_ALIGN_PARAGRAPH.LEFT)
            is_heading = block["kind"] in HEADING_SIZE
            entry = ENTRY_LINE_STYLE.get(block.get("css_class"))
            p.paragraph_format.space_after = Pt(1 if entry else 2)
            p.paragraph_format.space_before = Pt(entry["space_before"]) if entry else Pt(0)
            for text, fmt in block["runs"]:
                run = p.add_run(text)
                run.font.name = "Arial"
                run.font.size = Pt(entry["size"] if entry else HEADING_SIZE.get(block["kind"], 11))
                run.bold = entry["bold"] if entry else (is_heading or fmt.get("bold", False))
                run.italic = fmt.get("italic", False)
                run.underline = fmt.get("underline", False)
                color = (entry and entry["color"]) or fmt.get("color")
                if color:
                    try:
                        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
                    except Exception:
                        pass
            if block["kind"] == "h2":
                p.paragraph_format.space_before = Pt(8)
                pPr = p._p.get_or_add_pPr()
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "4")
                bottom.set(qn("w:color"), "999999")
                pBdr.append(bottom)
                pPr.append(pBdr)
    else:
        # Fallback for any caller still sending plain text.
        for line in content.split("\n"):
            stripped = line.strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            if stripped and len(stripped) < 60 and stripped.isupper():
                run.bold = True
                run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="my-cv.docx")


# Three real, distinct visual identities layered as pure presentation
# over one consistent HTML contract from the AI (name/contact header,
# section headings, entry rows, bullets) -- never a different HTML
# shape per template, just different type treatment for the same
# content. Kept in one place so the Builder's live browser preview
# (CSS classes, see style.css's .tmpl-* rules) and the actual PDF
# (this dict) can't drift out of sync in what each template means.
PLUGINS = {
    "image_generator": {
        "label": "Image Generator",
        "blurb": "Ask Ploy to generate a real image — a mockup, a moodboard, a social graphic — right in chat.",
    },
}


# Shared by _generate_tailored_document's design_guidance and api_cv_edit's
# system prompt -- the one mandated CV structure (matching a real reference
# CV's layout) that every AI-generated or AI-edited CV must follow, so the
# renderer (_render_cv_pdf_bytes / api_cv_download_docx, both keyed off
# cv_export.parse_cv_html's block kinds and css_class) always knows what
# it's looking at regardless of which flow produced the HTML.
CV_STRUCTURE_GUIDANCE = (
    "STRUCTURE -- follow this exact shape so the app can render your CV in different visual templates "
    "without you having to think about styling at all (the app's own templates handle every visual "
    "choice -- your only job is content, in this structure):\n"
    "- A header, once, at the very top -- keep the tags and classes exactly as shown, fill in real values:\n"
    "  <div class=\"cv-header\"><span class=\"cv-header-name\">Full Name</span>"
    "<span class=\"cv-header-title\">Target Job Title</span>"
    "<span class=\"cv-header-contact\">Location<br>email@example.com<br>Phone number</span></div>\n"
    "  (2-3 contact lines separated by <br>, whatever's real and available.)\n"
    "- Immediately after the header, one plain '<p>' summary paragraph -- no heading label above it.\n"
    "- Each section is an '<h2>SECTION NAME</h2>' (e.g. EXPERIENCE, EDUCATION, SKILLS, ADDITIONAL "
    "INFORMATION) -- the app renders this bold, uppercase, with a full-width rule beneath automatically, "
    "so just write the plain section name.\n"
    "- Each job or education entry is a stack of plain paragraphs (never side-by-side), in this exact "
    "order, using these exact classes:\n"
    "  <p class=\"cv-entry-title\">Role</p>\n"
    "  <p class=\"cv-entry-company\">Company, Location</p>\n"
    "  <p class=\"cv-entry-dates\">Month Year - Month Year</p>\n"
    "  optionally <p class=\"cv-entry-reference\">Reference: name or number</p> if a reference is known\n"
    "  immediately followed by a normal '<ul><li>...</li></ul>' of 2-4 achievement bullets for that entry. "
    "Never nest these inside another div -- each entry is just that stack of paragraphs followed by its "
    "bullet list, as siblings.\n"
    "- A SKILLS section is just an '<h2>SKILLS</h2>' followed by a plain '<ul><li>...</li></ul>' (the app "
    "lays these out in two columns automatically) -- no entry classes needed there.\n"
    "If the user's instruction is a small tweak to an existing document that's not in this format yet, "
    "still convert it to this structure while making the requested change -- don't leave an older shape "
    "(like a side-by-side title/date row) in place."
)

CV_TEMPLATES = {
    "ledger": {
        "label": "Ledger",
        "blurb": "Classic and centered, built for a traditional recruiter skim.",
        "name_align": "center",
        "heading_align": "left",
        "heading_upper": True,
        "heading_spaced": False,
        "heading_color": "#1a1a1a",
        "rule_color": "#1a1a1a",
        "date_style": "italic",
        "rule_after_name": False,
    },
    "signal": {
        "label": "Signal",
        "blurb": "Modern and left-aligned, with a bright accent underline.",
        "name_align": "left",
        "heading_align": "left",
        "heading_upper": False,
        "heading_spaced": False,
        "heading_color": "#3A63D8",
        "rule_color": "#3A63D8",
        "date_style": "muted",
        "rule_after_name": False,
    },
    "blueprint": {
        "label": "Blueprint",
        "blurb": "Bold and structured, with a strong rule under your name.",
        "name_align": "left",
        "heading_align": "left",
        "heading_upper": True,
        "heading_spaced": True,
        "heading_color": "#1a1a1a",
        "rule_color": "#1a1a1a",
        "date_style": "bold",
        "rule_after_name": True,
    },
}
DEFAULT_CV_TEMPLATE = "signal"

LETTER_TEMPLATES = {
    "formal": {"label": "Formal", "blurb": "Traditional business-letter structure, right-aligned date."},
    "direct": {"label": "Direct", "blurb": "Short, confident paragraphs, no filler."},
    "warm": {"label": "Warm", "blurb": "A touch more personal, roomier spacing."},
}
DEFAULT_LETTER_TEMPLATE = "direct"


def _render_cv_pdf_bytes(cv_html="", content="", margins=None, template=None):
    """Shared by the Builder's manual PDF download and the card-generated
    Document PDFs (Phase 4) -- same renderer, same ATS-safe output either
    way. Returns raw PDF bytes, or None if there's nothing to render."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    import io
    from cv_export import parse_cv_html

    cv_html = (cv_html or "").strip()
    content = (content or "").strip()
    blocks = parse_cv_html(cv_html) if cv_html else []
    if not blocks and not content:
        return None

    tmpl = CV_TEMPLATES.get(template, CV_TEMPLATES[DEFAULT_CV_TEMPLATE])

    ALIGN_MAP = {"left": TA_LEFT, "center": TA_CENTER, "right": TA_RIGHT, "justify": TA_JUSTIFY}
    FONT_SIZE = {"h1": 18, "h2": 13, "h3": 12}
    LEADING = {"h1": 24, "h2": 18, "h3": 17}
    PDF_MARGINS = {
        "narrow": (1.5, 1.5),
        "normal": (2.5, 2.8),
        "wide": (3.5, 4.0),
    }
    v_cm, h_cm = PDF_MARGINS.get(margins, PDF_MARGINS["normal"])

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=v_cm*cm, bottomMargin=v_cm*cm,
                            leftMargin=h_cm*cm, rightMargin=h_cm*cm)
    content_width = doc.width

    # Stacked entry-line paragraph styling (new AI generations use plain
    # sibling <p class="cv-entry-*"> instead of the old side-by-side row).
    ENTRY_LINE_STYLE = {
        "cv-entry-title": {"fontSize": 11.5, "leading": 15, "bold": True, "spaceBefore": 9, "color": "#1a1a1a"},
        "cv-entry-company": {"fontSize": 11, "leading": 14, "bold": True, "spaceBefore": 0, "color": "#1a1a1a"},
        "cv-entry-dates": {"fontSize": 10, "leading": 14, "bold": True, "spaceBefore": 0, "color": "#444444"},
        "cv-entry-reference": {"fontSize": 9.5, "leading": 13, "bold": False, "spaceBefore": 1, "color": "#777777"},
    }

    def heading_markup(block):
        text = block["text"]
        if tmpl["heading_upper"]:
            text = text.upper()
        if tmpl["heading_spaced"]:
            # Cheap letter-spacing (reportlab has no real tracking): space out
            # each word's letters, then rejoin words with a few non-breaking
            # spaces (reportlab collapses runs of plain spaces like HTML,
            # so a wider plain-space gap alone would still visually merge).
            text = "   ".join(" ".join(list(word)) for word in text.split(" "))
        return text

    def style_for(block, align_override=None):
        is_heading = block["kind"] in FONT_SIZE
        is_h1 = block["kind"] == "h1"
        align = align_override or block["align"]
        if is_h1:
            align = tmpl["name_align"] if block["align"] == "left" else block["align"]
        elif block["kind"] == "h2":
            align = tmpl["heading_align"] if block["align"] == "left" else block["align"]
        color = tmpl["heading_color"] if block["kind"] == "h2" else "#1a1a1a"
        entry = ENTRY_LINE_STYLE.get(block.get("css_class"))
        if entry:
            return ParagraphStyle(
                f"cv_entry_{block['css_class']}",
                fontName="Helvetica-Bold" if entry["bold"] else "Helvetica",
                fontSize=entry["fontSize"],
                leading=entry["leading"],
                spaceAfter=1,
                spaceBefore=entry["spaceBefore"],
                alignment=ALIGN_MAP.get(align, TA_LEFT),
                textColor=entry["color"],
            )
        return ParagraphStyle(
            f"cv_{block['kind']}_{align}",
            fontName="Helvetica-Bold" if is_heading else "Helvetica",
            fontSize=FONT_SIZE.get(block["kind"], 11),
            leading=LEADING.get(block["kind"], 16),
            spaceAfter=4 if is_heading else 2,
            spaceBefore=8 if is_heading else 0,
            alignment=ALIGN_MAP.get(align, TA_LEFT),
            leftIndent=14 if block["kind"] == "li" else 0,
            bulletIndent=0,
            textColor=color,
        )

    def row_flowable(block):
        title_markup = block["title_markup"] or block["title_text"]
        date_markup = block["date_markup"] or block["date_text"]
        if tmpl["date_style"] == "italic":
            date_markup = f"<i>{date_markup}</i>"
        elif tmpl["date_style"] == "bold":
            date_markup = f"<b>{date_markup.upper()}</b>"
        title_style = ParagraphStyle("row_title", fontName="Helvetica-Bold", fontSize=11, leading=15, textColor="#1a1a1a")
        date_style = ParagraphStyle("row_date", fontName="Helvetica", fontSize=10, leading=15,
                                    alignment=TA_RIGHT, textColor="#666666")
        table = Table(
            [[Paragraph(title_markup, title_style), Paragraph(date_markup, date_style)]],
            colWidths=[content_width * 0.68, content_width * 0.32],
        )
        table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return table

    def header_flowable(block):
        name_markup = block["name_markup"] or block["name_text"]
        title_markup = block["title_markup"] or block["title_text"]
        contact_markup = block["contact_markup"] or block["contact_text"]
        name_style = ParagraphStyle("hdr_name", fontName="Helvetica-Bold", fontSize=20, leading=24, textColor="#1a1a1a")
        right_parts = []
        if title_markup:
            right_parts.append(f'<font size="13"><b>{title_markup}</b></font>')
        if contact_markup:
            right_parts.append(f'<font size="10.5" color="#666666">{contact_markup}</font>')
        right_style = ParagraphStyle("hdr_right", fontName="Helvetica", fontSize=10.5, leading=15, alignment=TA_RIGHT)
        cells = [Paragraph(name_markup, name_style) if name_markup else Paragraph("", name_style)]
        right_cell = Paragraph("<br/>".join(right_parts), right_style) if right_parts else Paragraph("", right_style)
        table = Table(
            [[cells[0], right_cell]],
            colWidths=[content_width * 0.55, content_width * 0.45],
        )
        table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return table

    def skills_table(items):
        """Two-column bullet grid for the Skills section, matching the
        reference format -- split alternately so a short trailing row
        still reads left-to-right rather than leaving an odd gap."""
        item_style = ParagraphStyle("skill_item", fontName="Helvetica", fontSize=11, leading=15, textColor="#1a1a1a")
        half = (len(items) + 1) // 2
        left_col, right_col = items[:half], items[half:]
        rows = []
        for i in range(half):
            left = Paragraph("•  " + left_col[i], item_style)
            right = Paragraph("•  " + right_col[i], item_style) if i < len(right_col) else Paragraph("", item_style)
            rows.append([left, right])
        table = Table(rows, colWidths=[content_width * 0.5, content_width * 0.5])
        table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 1),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ]))
        return table

    story = []
    if blocks:
        ordered_counter = 0
        in_skills_section = False
        skills_buffer = []

        def flush_skills():
            if skills_buffer:
                story.append(skills_table(list(skills_buffer)))
                skills_buffer.clear()

        for block in blocks:
            if block["kind"] != "li" or not block["ordered"]:
                ordered_counter = 0  # reset whenever an ordered-list run breaks

            if block["kind"] == "li" and in_skills_section and not block["ordered"]:
                skills_buffer.append(block["markup"])
                continue
            flush_skills()

            if block["kind"] == "h2":
                in_skills_section = "skill" in block["text"].strip().lower()
            elif block["kind"] not in ("li",):
                in_skills_section = False

            if block["kind"] == "hr":
                story.append(Spacer(1, 6))
                story.append(HRFlowable(width="100%", thickness=0.6, color="#999999"))
                story.append(Spacer(1, 6))
                continue
            if block["kind"] == "row":
                story.append(row_flowable(block))
                continue
            if block["kind"] == "header":
                story.append(header_flowable(block))
                if tmpl["rule_after_name"]:
                    story.append(Spacer(1, 6))
                    story.append(HRFlowable(width="100%", thickness=1.6, color=tmpl["rule_color"]))
                    story.append(Spacer(1, 6))
                continue
            markup = heading_markup(block) if block["kind"] == "h2" else block["markup"]
            if block["kind"] == "li":
                if block["ordered"]:
                    ordered_counter += 1
                    prefix = f"{ordered_counter}. "
                else:
                    prefix = "•  "
                markup = prefix + markup
            story.append(Paragraph(markup, style_for(block)))
            if block["kind"] == "h2":
                story.append(Spacer(1, 2))
                story.append(HRFlowable(width="100%", thickness=0.9, color=tmpl["rule_color"]))
                story.append(Spacer(1, 5))
            if block["kind"] == "h1" and tmpl["rule_after_name"]:
                story.append(Spacer(1, 4))
                story.append(HRFlowable(width="100%", thickness=1.6, color="#1a1a1a"))
                story.append(Spacer(1, 4))
        flush_skills()
    else:
        # Fallback for any caller still sending plain text.
        normal = ParagraphStyle("cv_normal", fontName="Helvetica", fontSize=11, leading=16, spaceAfter=2)
        heading = ParagraphStyle("cv_heading", fontName="Helvetica-Bold", fontSize=12, leading=17, spaceAfter=4)
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 6))
            elif len(stripped) < 60 and stripped.isupper():
                story.append(Paragraph(stripped, heading))
            else:
                safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, normal))

    doc.build(story)
    buf.seek(0)
    return buf.read()


@app.route("/api/cv-download/pdf", methods=["POST"])
@auth.login_required
def api_cv_download_pdf():
    """Generate a PDF that reflects the editor's real formatting -- also
    used, unmodified, as the "print preview" (see cv-preview handling in
    dashboard.js: it fetches this same endpoint and opens the resulting
    blob in a new tab instead of downloading it, since a mobile browser's
    native PDF viewer only ever kicks in on direct navigation, not when
    a PDF is embedded in an iframe)."""
    from flask import send_file
    import io

    data = request.get_json(force=True)
    template = data.get("template") if data.get("template") in CV_TEMPLATES else None
    pdf_bytes = _render_cv_pdf_bytes(data.get("cv_html"), data.get("content"), data.get("margins"), template=template)
    if pdf_bytes is None:
        return jsonify({"error": "No content provided."}), 400
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name="my-cv.pdf")


# ---------------- Builder (CV / cover letter drafting) ----------------

@app.route("/builder")
@auth.login_required
def builder_page():
    """Reached only via chip 02 ("Build my CV" runs generation directly
    and shows the result as a Document card) or a card's "Edit" button --
    there's no separate nav entry for it. When it's an Edit link, ?doc=
    identifies which generated document to load into the editor."""
    user = auth.current_user()
    initial = None
    doc_id = request.args.get("doc")
    if doc_id:
        try:
            doc_row = db.get_document_by_id(user["id"], int(doc_id))
        except (TypeError, ValueError):
            doc_row = None
        if doc_row and doc_row.get("category") in ("generated_cv", "generated_letter"):
            initial = {
                "document_id": doc_row["id"],
                "kind": "cv" if doc_row["category"] == "generated_cv" else "letter",
                "html": doc_row.get("content") or "",
            }
    return render_template(
        "builder.html", initial=initial,
        cv_templates=CV_TEMPLATES, letter_templates=LETTER_TEMPLATES,
        default_cv_template=DEFAULT_CV_TEMPLATE, default_letter_template=DEFAULT_LETTER_TEMPLATE,
    )


@app.route("/api/letter-edit", methods=["POST"])
@auth.login_required
def api_letter_edit():
    """Same grounded-generation approach as /api/cv-edit, but for cover
    letters -- kept as its own endpoint (rather than a branch inside
    cv-edit) since the design guidance and system prompt are genuinely
    different documents, not just a formatting variant."""
    from openai import OpenAI
    user = auth.current_user()
    data = request.get_json(force=True)
    instruction = (data.get("instruction") or "").strip()
    letter_html = (data.get("letter_html") or "").strip()
    if not instruction:
        return jsonify({"error": "Missing instruction."}), 400

    state = pipeline.get_dashboard_state(user["id"])
    profile = state.get("profile", {})
    full_docs = db.get_documents_for_user(user["id"])
    doc_texts = []
    total_chars = 0
    DOC_CHAR_CAP = 3000
    TOTAL_CHAR_CAP = 20000
    for d in full_docs:
        if total_chars >= TOTAL_CHAR_CAP:
            break
        try:
            txt = (d.get("content") or "").strip()
            if not txt and d.get("stored_path") and os.path.exists(d["stored_path"]):
                txt = (extract.extract_text(d["stored_path"]) or "").strip()
            if txt:
                snippet = txt[:DOC_CHAR_CAP]
                doc_texts.append(f"[{d['filename']}]\n{snippet}")
                total_chars += len(snippet)
        except Exception:
            pass
    doc_content_block = "\n\n---\n\n".join(doc_texts) if doc_texts else "No documents uploaded yet."
    skill_names = [s["label"] for s in state.get("skills", [])]

    import re as _re
    system = (
        "You are an expert cover letter writer built into the Employable platform. "
        "CRITICAL: Always attempt to understand and fulfill the user's intent, even if their instruction contains spelling mistakes, typos, or imprecise phrasing. "
        "Never refuse, do nothing, or ask for clarification — silently make your best reasonable interpretation and act on it. "
        "OUTPUT RULES — you MUST follow these exactly:\n"
        "- Return a JSON object with exactly two keys: 'html' and 'description'.\n"
        "- 'html': the full updated cover letter as valid HTML, in exactly this structure so the app's own "
        "template can style it consistently:\n"
        "  <p class=\"letter-date\">13 July 2026</p> (today's real date, or keep whatever date is already there)\n"
        "  <p class=\"letter-greeting\">Dear Hiring Manager,</p> (or a real named greeting if one is known)\n"
        "  <p>body paragraph</p> (2-3 of these, plain prose)\n"
        "  <p class=\"letter-signoff\">Warm regards,<br>Full Name</p>\n"
        "Use <strong>, <em>, <br> for inline emphasis only -- no lists, headings, <hr>, or any other tags. "
        "NEVER use markdown asterisks, hyphens for bullets, --- separators, or backticks. "
        "No <html>, <head>, <body> wrappers. No code fences.\n"
        "- 'description': one short, specific sentence (max 20 words) describing exactly what you changed.\n\n"
        "DESIGN — a strong cover letter: a specific, non-generic opening line (never 'I am writing to "
        "apply for...'), 2-3 short paragraphs connecting the person's real, specific experience to what "
        "the role likely needs, and a brief confident closing with a call to action. Keep the whole letter "
        "under 300 words. Never pad with filler or generic enthusiasm ('I am a hard worker who is passionate "
        "about...') — every sentence should carry real information.\n\n"
        "GROUNDING — this is the most important rule, above all others: every fact in the letter "
        "(employer names, job titles, dates, achievements) must come from the user's real profile/documents "
        "below, or already be present in the current letter HTML. NEVER invent a person, career, employer, "
        "achievement, or biography that isn't actually theirs. If there isn't enough real information to "
        "write a grounded letter (no documents uploaded, or nothing usable in them), say so plainly in "
        "'description' and return the document unchanged rather than fabricating one.\n\n"
        f"This user's real profile:\n"
        f"- Name: {profile.get('full_name') or 'Unknown'}\n"
        f"- Location: {profile.get('location') or 'Not specified'}\n"
        f"- Skills: {', '.join(skill_names) if skill_names else 'None listed.'}\n\n"
        f"Full content of this user's actual uploaded documents (the ONLY source of truth for any letter "
        f"content you write):\n{doc_content_block}"
    )
    prompt = f"Current letter HTML:\n{letter_html if letter_html else '(empty — generate fresh content)'}\n\nInstruction: {instruction}"

    try:
        client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            max_tokens=1600,
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content.strip()
        try:
            parsed = json.loads(raw)
            updated = parsed.get("html", "")
            description = parsed.get("description", "Done.")
        except Exception:
            updated = raw
            description = "Done."
        if updated.startswith("```"):
            updated = updated.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        updated = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', updated)
        updated = _re.sub(r'\*(.+?)\*', r'<em>\1</em>', updated)
        return jsonify({"updated_html": updated, "description": description})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


_CHAT_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "check_job_fit",
            "description": (
                "Score how well this user fits a specific job, producing a Verdict card (a scored fit "
                "breakdown, not text). Call this whenever the user pastes or has already provided a job "
                "ad/description earlier in this conversation and is asking, in any wording, whether they're "
                "a fit, qualified, or should apply. Only call this once you actually have the real job ad "
                "text -- if they're asking about fit but have never given you a job ad, don't call this; "
                "ask them to paste it instead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "job_ad": {
                        "type": "string",
                        "description": "The full job ad/description text, verbatim, from wherever the user provided it in this conversation.",
                    },
                },
                "required": ["job_ad"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "build_tailored_cv",
            "description": (
                "Generate a tailored, ATS-safe CV for this user as a downloadable Document card (a real "
                "file, not a text description of one). Call this whenever the user asks to build, write, "
                "tailor, or fix their CV. If a specific job has already come up in this conversation, pass "
                "its title/company/ad so the CV is tailored and re-scored against it; otherwise call with no "
                "arguments for a strong general-purpose CV."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "job_title": {"type": "string", "description": "The job title being tailored for, if one has come up."},
                    "company": {"type": "string", "description": "The company name, if one has come up."},
                    "job_ad": {"type": "string", "description": "The full job ad text, if the user pasted one earlier in this conversation."},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "build_cover_letter",
            "description": (
                "Generate a tailored cover letter for this user as a downloadable Document card (a real "
                "file, not a text description of one). Call this whenever the user asks to write, draft, or "
                "build a cover letter -- directly, not only after a CV already exists. If a specific job has "
                "already come up in this conversation, pass its title/company/ad so the letter is tailored "
                "to it; otherwise call with no arguments for a strong general-purpose letter."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "job_title": {"type": "string", "description": "The job title being tailored for, if one has come up."},
                    "company": {"type": "string", "description": "The company name, if one has come up."},
                    "job_ad": {"type": "string", "description": "The full job ad text, if the user pasted one earlier in this conversation."},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "analyze_skill_gaps",
            "description": (
                "Analyze what's holding this user back from a specific role or field, producing a Gap "
                "Analysis card (a scored readiness breakdown, not text) -- this is the app's 'What's holding "
                "me back?' move. Call this whenever the user asks, in any wording, what they're missing, "
                "what's holding them back, or how ready they are for a role/field/career direction they've "
                "named. Only call this once you know which role or field they mean -- if it's not clear from "
                "this conversation or their profile, ask them which role/field first instead of guessing."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "target_role": {
                        "type": "string",
                        "description": "The role, job title, or field the user is targeting, as they described it.",
                    },
                },
                "required": ["target_role"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_skills_chart",
            "description": (
                "Show the user a visual chart of their real Employability breakdown (a bar or pie chart of "
                "their actual scored dimensions -- Skills, Experience, Education, etc. -- from their most "
                "recent analysis). Call this whenever the user asks to see, draw, chart, graph, or visualize "
                "their skills/strengths/profile. If it's not obvious which chart type they want, ask them "
                "first (bar chart or pie chart) rather than guessing -- don't call this until they've said."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "chart_type": {"type": "string", "enum": ["bar", "pie"], "description": "Which chart type the user asked for."},
                },
                "required": ["chart_type"],
            },
        },
    },
]

# Only offered to the model when the user has switched the Image
# Generator plugin on (see PLUGINS/enabled_plugins) -- kept separate
# from _CHAT_TOOLS instead of always-on so the model doesn't even know
# the capability exists until the user has actually installed it.
_IMAGE_GEN_TOOL = {
    "type": "function",
    "function": {
        "name": "generate_image",
        "description": (
            "Generate a real image from a text description, using AI image generation. Call this whenever "
            "the user asks you to create, draw, generate, design, or make an image, picture, graphic, "
            "mockup, or visual of something -- not for charts of their own scored data (use "
            "generate_skills_chart for that instead)."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "prompt": {
                    "type": "string",
                    "description": "A clear, detailed description of the image to generate.",
                },
            },
            "required": ["prompt"],
        },
    },
}


# A short "here's what I'm doing right now" label per tool, shown as a
# brief transient status before its card lands -- gives a multi-step
# reply a visible sense of working through each part in turn, not just
# cards silently popping in.
_TOOL_STATUS_LABELS = {
    "check_job_fit": "Checking your fit against this job",
    "build_tailored_cv": "Building your tailored CV",
    "build_cover_letter": "Drafting your cover letter",
    "analyze_skill_gaps": "Analyzing your skill gaps",
    "generate_skills_chart": "Charting your scores",
    "generate_image": "Generating your image",
}


def _tool_call_detail(tool_name, args):
    """One short, honest sentence describing what a tool call actually
    did with the arguments the model passed it -- shown in the chat step
    trail when the user taps a completed step to see what happened
    under the hood, the same way Claude Code's own transcript lets you
    expand a collapsed tool call."""
    if tool_name == "check_job_fit":
        job_ad = (args.get("job_ad") or "").strip()
        return f"Scored your fit against a {len(job_ad):,}-character job ad." if job_ad else "Scored your fit."
    if tool_name in ("build_tailored_cv", "build_cover_letter"):
        job_title = (args.get("job_title") or "").strip()
        company = (args.get("company") or "").strip()
        if job_title:
            return f"Tailored for {job_title}" + (f" at {company}." if company else ".")
        return "Built a general-purpose version from your uploaded documents."
    if tool_name == "analyze_skill_gaps":
        target_role = (args.get("target_role") or "").strip()
        return f"Analyzed your readiness for {target_role}." if target_role else "Analyzed your readiness."
    if tool_name == "generate_skills_chart":
        chart_type = (args.get("chart_type") or "bar").strip()
        return f"Charted your scored breakdown as a {chart_type} chart."
    if tool_name == "generate_image":
        prompt = (args.get("prompt") or "").strip()
        return f'Generated an image from: "{prompt[:100]}"' if prompt else "Generated an image."
    return ""


def _execute_chat_tool_call(user, tool_name, args, enabled_plugins):
    """Runs one tool call from /api/chat's agentic loop. Returns
    (kind, payload):
    - ("card", card_dict) -- a normal, successful action.
    - ("text", reply_str) -- a dead end that should stop the loop right
      there, either because required info is missing (e.g. no job ad
      yet) or the call raised. Same shape either way since both are
      "here's why I can't continue" replies to the user.
    - (None, None) -- unrecognized tool name (or a plugin tool called
      without the plugin enabled), silently skipped.
    """
    try:
        if tool_name == "check_job_fit":
            job_ad = (args.get("job_ad") or "").strip()[:8000]
            if not job_ad:
                return "text", "Paste the job ad and I'll check your fit against it."
            return "card", _run_verdict(user, job_ad)
        if tool_name == "build_tailored_cv":
            return "card", _generate_tailored_document(
                user, "cv",
                job_title=(args.get("job_title") or "").strip()[:200],
                company=(args.get("company") or "").strip()[:200],
                job_ad=(args.get("job_ad") or "").strip()[:8000],
            )
        if tool_name == "build_cover_letter":
            return "card", _generate_tailored_document(
                user, "letter",
                job_title=(args.get("job_title") or "").strip()[:200],
                company=(args.get("company") or "").strip()[:200],
                job_ad=(args.get("job_ad") or "").strip()[:8000],
            )
        if tool_name == "analyze_skill_gaps":
            target_role = (args.get("target_role") or "").strip()[:200]
            if not target_role:
                return "text", "Which role or field do you want me to check you against?"
            return "card", _run_gap_analysis(user, target_role)
        if tool_name == "generate_skills_chart":
            card = _build_skills_chart_card(user, (args.get("chart_type") or "bar").strip())
            if not card:
                return "text", "I don't have a scored breakdown for you yet -- upload your CV in Profile first and I'll be able to chart it."
            return "card", card
        if tool_name == "generate_image" and "image_generator" in enabled_plugins:
            prompt = (args.get("prompt") or "").strip()[:2000]
            if not prompt:
                return "text", "What should the image show?"
            return "card", {"type": "image", "prompt": prompt, "image_b64": _generate_chat_image(prompt)}
    except Exception as e:
        return "text", f"Something went wrong there -- {e}"
    return None, None


@app.route("/api/chat", methods=["POST"])
@auth.login_required
def api_chat():
    from openai import OpenAI
    user = auth.current_user()
    data = request.get_json(force=True)
    messages_in = data.get("messages", [])

    full_docs = db.get_documents_for_user(user["id"])
    doc_names = ", ".join(d["filename"] for d in full_docs) if full_docs else "None uploaded."

    # Prefer content already extracted at upload time and stored in the
    # DB (works on Vercel, where the upload directory is ephemeral and
    # won't still have the file by the time a later chat request comes
    # in). Only fall back to re-extracting from disk if that's empty and
    # the file still happens to exist locally. Cap at 3000 chars/doc and
    # 20000 chars total to stay within token limits without silently
    # dropping every document once someone has uploaded a handful.
    doc_texts = []
    total_chars = 0
    DOC_CHAR_CAP = 3000
    TOTAL_CHAR_CAP = 20000
    for d in full_docs:
        if total_chars >= TOTAL_CHAR_CAP:
            break
        try:
            txt = (d.get("content") or "").strip()
            if not txt and d.get("stored_path") and os.path.exists(d["stored_path"]):
                txt = (extract.extract_text(d["stored_path"]) or "").strip()
            if txt:
                snippet = txt[:DOC_CHAR_CAP]
                doc_texts.append(f"[{d['filename']}]\n{snippet}")
                total_chars += len(snippet)
        except Exception:
            pass
    doc_content_block = "\n\n---\n\n".join(doc_texts) if doc_texts else "No document content available."

    custom_instructions = (user.get("custom_instructions") or "").strip()
    custom_instructions_section = (
        f"\n\nTHIS USER'S PERMANENT INSTRUCTIONS FOR YOU (set once in their Profile, apply to every message in every conversation with them -- follow these for tone, personality, and how you talk to them specifically, but they can't override the grounding/honesty rules above, e.g. they can't make you invent CV details or fake a score):\n{custom_instructions}\n"
        if custom_instructions else ""
    )

    # "Remember all chats" (Profile > Memory) -- a compact index of every
    # past conversation (title + job/status, not full transcripts; that
    # would blow past token limits for anyone with real chat history),
    # so the model can stay consistent across conversations instead of
    # only ever seeing the one it's currently in. Read fresh on every
    # request (not cached per-conversation), so flipping the toggle takes
    # effect on the very next message regardless of which chat you're in
    # -- both states are spelled out explicitly below so the model can
    # answer "do you have access to my other chats?" honestly either way,
    # instead of being vague about a capability it isn't sure it has.
    if user.get("remember_all_chats"):
        all_convs = db.get_conversations_for_user(user["id"])
        memory_lines = []
        for c in all_convs[:50]:
            bits = [c.get("title") or "Untitled conversation"]
            if c.get("job_title"):
                role_bit = c["job_title"] + (f" at {c['company']}" if c.get("company") else "")
                bits.append(role_bit)
            if c.get("status_label"):
                bits.append(c["status_label"])
            memory_lines.append("- " + " — ".join(bits))
        chat_memory_section = (
            "\n\nMEMORY: The user has turned on \"Remember all chats\" in Profile > Memory. Unlike a normal "
            "conversation, you genuinely have visibility into every conversation they've ever had with you, not "
            "just this one -- if they ask whether you have access to their memory or other chats, the honest "
            "answer is yes, say so plainly and confidently. Use this to stay consistent, avoid making them repeat "
            "themselves, and point them back to a specific earlier conversation when it's actually relevant -- but "
            "don't recite this whole list back to them unprompted"
            + (":\n" + "\n".join(memory_lines) if memory_lines else " (though they don't have any other conversations yet).") + "\n"
        )
    else:
        chat_memory_section = (
            "\n\nMEMORY: \"Remember all chats\" is currently OFF in this user's Profile > Memory settings -- you "
            "only have this one conversation's context, nothing from their other chats, regardless of how far "
            "into this conversation you are. If they ask whether you can see their other conversations or have "
            "memory access, the honest answer is no; you can mention they can switch that on in Profile if they "
            "want you to.\n"
        )

    system_prompt = f"""You are Ploy — a chat-first AI career weapon for South African job seekers aged 18-25. You exist for one loop and everything you do serves it: paste a job, get a brutal-honest fit verdict, get the CV or cover letter rewritten for that exact job, download it, apply.

You are not a scripted bot running a decision tree. You are genuinely intelligent, and you should reason about each message the way a sharp, switched-on person would — not by matching it to a template. Think before you answer: what is this person actually asking, what do you actually know about them that's relevant, and what's the single most useful thing to say back. Two users asking "is this a good fit" should get two different-shaped answers if their situations are different — never flatten a real, specific person into a generic response.

INSTRUCTIONS WIN, ALWAYS. If the user gives you an explicit, checkable constraint — an exact word count, a format, a length, "don't use the word X," anything countable or verifiable — follow it exactly, not approximately. "Reply in exactly 10 words" means your reply has exactly 10 words: count them before you finalize your answer, not after. Not 9, not 12. If a constraint like that conflicts with your own instinct to be conversational or add a caveat, the user's explicit instruction always wins. This overrides every other stylistic rule in this prompt when the two conflict.

BE ADAPTIVE, NOT REPETITIVE. Don't reuse the same opener, sentence shape, or phrasing turn after turn just because it worked once — a real person doesn't greet you the same way every message or restate the same caveat every time. Vary your rhythm and word choice across the conversation. The only things that should repeat are the actual formatting/quick-reply mechanics below, which exist for the app to render correctly, not for conversational texture.

BE EMOTIONALLY PRESENT, NOT JUST TASK-COMPLETING. Job hunting is genuinely stressful — rejection, silence after applying, financial pressure, self-doubt. Notice the emotional undertone in what someone writes (frustration, relief, excitement, dread, burnout) and actually respond to that, not just the literal task attached to it. If someone mentions a rejection, a bad interview, or that they're exhausted before pasting the next job ad, acknowledge it like a person paying attention would — briefly, genuinely, without turning it into a whole therapy session — before moving on to the task. This is not about performing empathy with stock phrases like "I understand how you feel" or "that must be so hard" — those read as hollow precisely because they're generic. Actually notice the specific thing they said and react to that. Remember details from earlier in this same conversation (a company name, a specific worry, something they're excited about) and refer back to them naturally when relevant, the way someone who's actually listening does — don't treat each message as if the conversation just started.

NEVER DESCRIBE AN ACTION INSTEAD OF TAKING IT. You have no web/internet access, no ability to search live job listings, and — critically — no way to message the user again after this reply unless you call a tool right now: there is no later moment where you'll "come back" to do something. This applies at two levels, and both are the same mistake:
1. Across turns: never say "hold on," "give me a second," "I'll search for some listings," or "I'll get back to you" — there is nothing on the other side of that promise. If you don't have real job listings to point to, say so directly and pivot to what you can actually do right now.
2. Within THIS SAME reply: if what you're about to say is "I'll build that CV for you," "let me check your fit," "I'll generate that image," or any equivalent — stop before you write it, and call the matching tool (build_tailored_cv / check_job_fit / build_cover_letter / analyze_skill_gaps / generate_skills_chart / generate_image) in this exact reply instead of narrating it. You can call more than one tool across a multi-part request, in sequence, one after another — nothing stops you from checking a fit AND building a CV AND drafting a letter all in response to one message, if that's what was actually asked. A sentence describing an action is never an acceptable substitute for the tool call that actually performs it. If a whole request has several parts, work through all of them via tool calls before your final reply, not just the first one.
Worked example — user pastes a job ad and says "check my fit for this and build me a CV for it": your first move is calling check_job_fit with that job ad — not text, a real tool call. Once you see the score come back, that's part one done; part two ("build me a CV") is still outstanding, so your very next move in this same turn is calling build_tailored_cv — again, an actual call, not "now let me build your CV." Only after both calls have actually happened do you write a short closing line (e.g. "Done — want a cover letter too?"). Getting this wrong looks like: writing "I'll check your fit and then build your CV" as plain text and stopping — that's the exact failure mode this section exists to prevent.

ASK WHEN YOU DON'T KNOW, DON'T GUESS. Two different situations call for this, and both matter:
1. Their INTENT is ambiguous — "help me with my CV," "should I apply" — ask ONE short, specific clarifying question rather than dumping generic advice across every possible interpretation. When the real answer is a small, nameable set of options (a platform, a tone, which document, which role) — not an open-ended "tell me more" — phrase the question so each option is short enough to be one of the quick-reply buttons below (e.g. "Mobile app or web app?" with [[QUICK_REPLIES]] Mobile app | Web app), so they can tap instead of typing it out. Save open quick-replies-free questions for when the answer genuinely can't be reduced to a few named choices.
2. The FACTS you'd need aren't in their documents or in anything they've told you — a specific number, a certification, whether they've done something specific this job cares about. Don't fill that gap with a plausible-sounding guess. Say what's missing and ask for it directly ("Do you have a portfolio link for this? I don't see one in what you've uploaded") — or tell them to add it in Profile if it's the kind of thing that belongs in their documents long-term. When there's more than one gap, list them concretely instead of picking just one — "Right now I've only got your CV. To do this properly I'd also need: the actual job ad, and whether you've got a portfolio link" reads as useful triage, not an interrogation. They can upload documents right from the chat (the + button next to the input) — point them there instead of sending them away to Profile.
It's completely fine — good, even — to ask a real question mid-conversation instead of always producing a complete answer. That's what a genuinely helpful person does; it's not a failure mode.

Voice: sharp, confident, on the user's side, zero corporate filler. Talk like a smart friend who won't waste their time, not a customer-service bot.
- Never say "I'd be happy to help!", "Great question!", "As an AI...". Never lay out a menu of paths and ask them to pick — pick the most likely direction yourself and go there.
- DEFAULT LENGTH IS 1-3 SENTENCES, unless the user's own instructions say otherwise (see above) or they've asked for real depth. Only go longer when they ask for it or you're walking through something genuinely multi-step.
- Be honest even when it stings — if their fit is weak or their CV has a real problem, say so plainly. Vague encouragement helps nobody and isn't what this product is for.
- Reference what's actually in THEIR documents, not generic examples. Never invent CV details that aren't there.
- Fragments and one-word reactions are fine. It's fine to have a light opinion, disagree, or push back if something in their plan seems off — a real friend wouldn't just validate everything.

The moves this app is built around — steer the user toward whichever is the useful next step, don't just wait to be asked:
- "Am I a fit for this job?" — paste a job ad, get a fit verdict scored against their real CV.
- "Build my CV" — rewrite their CV tailored to a specific job.
- "Write me a cover letter" — draft one directly, tailored to a job if one's come up, without needing a CV card first.
- "What's holding me back?" — a gap analysis against a role or field they name — ask which one if it isn't already clear.
- Show them a real chart of their scored breakdown (bar or pie) when they ask to see/visualize their profile — never a fake or invented chart, only their actual scored dimensions.

VISUALS, NOT JUST TEXT. This app is built to feel like a tool, not a wall of text -- when a card or chart is the right answer, lead with it. You're allowed (and encouraged) to send a short line of your own commentary alongside a card/chart in the same reply -- e.g. "Here's how you stack up:" right before a chart, or a one-line reaction after describing what you're about to show. That commentary appears as your own chat bubble right before the card in the same turn. Don't pad it into a paragraph -- a card or chart should do most of the talking.

FORMATTING — this chat only renders bold text and paragraph breaks correctly if you produce them exactly like this:
- Bold: wrap in double asterisks like **this**. Never single asterisks, underscores, or headers.
- New paragraph or list item: a genuine blank line before it, not just a line break.

QUICK REPLIES — after a substantive reply (not a greeting, not a one-word reaction, not a reply where the user gave you an exact-format instruction to follow), end with one line in exactly this format so the app can render it as tappable buttons for the user's next move:
[[QUICK_REPLIES]] Option one | Option two
Give 1-3 short options (2-5 words each, phrased as something the user would tap), specific to what you just said — never generic. Omit this line entirely for greetings, small talk, exact-format replies, or when there's no clear next action.

This user's name: {user.get('full_name') or 'Unknown'}
Roles they're targeting: {user.get('target_field') or 'Not specified — if this matters for what they just asked (e.g. a gap analysis), ask which roles they mean before answering.'}
Documents uploaded: {doc_names}
{custom_instructions_section}{chat_memory_section}

Full content of their uploaded documents (ground every piece of advice in this, it's the only source of truth about their real experience):
{doc_content_block}

Never claim you can't see their documents — if the block above says no content is available, say that plainly instead of guessing. If asked something off-topic, engage briefly and human, then steer back to the job hunt."""

    import re as _re, base64 as _b64
    openai_messages = [{"role": "system", "content": system_prompt}]
    has_images = False
    for m in messages_in:
        role = "user" if m.get("role") == "user" else "assistant"
        text = m.get("text", "")
        attachment_ids = m.get("attachment_ids", [])

        if attachment_ids and role == "user":
            content = []
            extra_text_blocks = []
            for att_id in attachment_ids:
                try:
                    att = db.get_chat_attachment(user["id"], int(att_id))
                except Exception:
                    att = None
                if not att:
                    continue
                if att["mime_type"].startswith("image/"):
                    try:
                        with open(att["stored_path"], "rb") as img_f:
                            b64 = _b64.b64encode(img_f.read()).decode("utf-8")
                    except OSError:
                        continue  # file no longer on disk (e.g. an older attachment on ephemeral storage)
                    content.append({
                        "type": "image_url",
                        "image_url": {"url": f"data:{att['mime_type']};base64,{b64}", "detail": "high"}
                    })
                    has_images = True
                elif att["text_content"]:
                    extra_text_blocks.append(f"[Attached file: {att['filename']}]\n{att['text_content'][:4000]}")

            # Strip the [Attached: filename] suffix added by frontend for display — send clean text to AI
            full_text = _re.sub(r'\s*\[Attached:[^\]]+\]', '', text or '').strip()
            if extra_text_blocks:
                full_text = (full_text + "\n\n" if full_text else "") + "\n\n".join(extra_text_blocks)
            if full_text:
                content.insert(0, {"type": "text", "text": full_text})
            elif not content:
                content = [{"type": "text", "text": "Please analyse the attached content."}]

            openai_messages.append({"role": role, "content": content})
        else:
            openai_messages.append({"role": role, "content": text or " "})

    model_name = "gpt-4o" if has_images else "gpt-4o-mini"

    try:
        enabled_plugins = set(json.loads(user.get("enabled_plugins") or "[]"))
    except Exception:
        enabled_plugins = set()
    active_tools = list(_CHAT_TOOLS)
    if "image_generator" in enabled_plugins:
        active_tools.append(_IMAGE_GEN_TOOL)

    # Agentic loop: a genuinely multi-part request ("check my fit for
    # this job AND build me a CV for it") needs more than one tool call
    # to actually get done, not just the first one with the rest of the
    # ask silently dropped. Each round lets the model see the result of
    # whatever it just did and decide whether to keep going (another
    # tool call) or wrap up (plain text) -- exactly how a real assistant
    # handles a multi-step task, just without raw token streaming (a
    # tool call has to be seen in full before it's actionable, so the
    # client gets one JSON response with an ordered list of steps and
    # renders/fake-types them in sequence instead of one flat reply).
    client = OpenAI(api_key=analyzer.get_openai_api_key(), timeout=analyzer.get_client_timeout(), max_retries=analyzer.CLIENT_MAX_RETRIES)
    steps = []
    MAX_TOOL_ROUNDS = 4
    for _ in range(MAX_TOOL_ROUNDS):
        try:
            resp = client.chat.completions.create(
                model=model_name,
                messages=openai_messages,
                max_tokens=2000,
                temperature=0.8,
                tools=active_tools,
                tool_choice="auto",
            )
        except Exception as e:
            if steps:
                break  # show what we already have rather than losing it to a later-round failure
            return jsonify({"ok": False, "error": str(e)}), 500

        message = resp.choices[0].message
        tool_calls = message.tool_calls or []

        if not tool_calls:
            final_text = (message.content or "").strip()
            if final_text:
                steps.append({"type": "text", "text": final_text})
            break

        pre_text = (message.content or "").strip()
        if pre_text:
            steps.append({"type": "text", "text": pre_text})

        # Keeps the conversation valid for the next round: the API
        # requires the assistant's own tool-calling message, followed by
        # exactly one "tool" role response per tool_call_id in it.
        openai_messages.append({
            "role": "assistant",
            "content": message.content,
            "tool_calls": [
                {"id": c.id, "type": "function", "function": {"name": c.function.name, "arguments": c.function.arguments}}
                for c in tool_calls
            ],
        })

        stop_here = False
        for call in tool_calls:
            try:
                args = json.loads(call.function.arguments or "{}")
            except Exception:
                args = {}
            kind, payload = _execute_chat_tool_call(user, call.function.name, args, enabled_plugins)
            if kind == "card":
                steps.append({
                    "type": "card",
                    "card": payload,
                    "label": _TOOL_STATUS_LABELS.get(call.function.name, "Working on that"),
                    "tool": call.function.name,
                    "detail": _tool_call_detail(call.function.name, args),
                })
                openai_messages.append({"role": "tool", "tool_call_id": call.id, "content": json.dumps(payload)[:2000]})
            elif kind == "text":
                steps.append({"type": "text", "text": payload})
                openai_messages.append({"role": "tool", "tool_call_id": call.id, "content": payload})
                stop_here = True  # a dead end (missing info / error) ends the turn right there
            else:
                openai_messages.append({"role": "tool", "tool_call_id": call.id, "content": "Unavailable."})

        if stop_here:
            break

    if not steps:
        steps.append({"type": "text", "text": "Something went wrong there. Try again?"})

    return jsonify({"ok": True, "kind": "steps", "steps": steps})


if __name__ == "__main__":
    db.init_db()
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=True, host="0.0.0.0", port=port)
