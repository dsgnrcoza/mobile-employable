"""
extract.py
----------
(Renamed from doc_extract.py)

Pulls plain text out of whatever the user uploaded, regardless of format,
so it can be handed to the OpenAI API as a single block of text.

Supported: .pdf, .docx, .doc (best-effort), .txt, .jpg/.jpeg/.png/.tiff/.tif (OCR)

Install:
    pip install pypdf python-docx pillow pytesseract
    # pytesseract also needs the Tesseract OCR binary itself installed on
    # the OS (not just the pip package):
    #   Windows: https://github.com/UB-Mannheim/tesseract/wiki
    #   macOS:   brew install tesseract
    #   Linux:   sudo apt install tesseract-ocr
"""

import os
import shutil
import sys

# pytesseract requires a system-level Tesseract binary.
# On Vercel and other serverless runtimes it's unavailable — import
# gracefully so the rest of the app still works for PDF/DOCX/TXT.
_pytesseract = None
try:
    import pytesseract as _pytesseract
    _tesseract_on_path = shutil.which("tesseract")
    if _tesseract_on_path:
        _pytesseract.pytesseract.tesseract_cmd = _tesseract_on_path
    elif sys.platform.startswith("win"):
        _default_windows_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(_default_windows_path):
            _pytesseract.pytesseract.tesseract_cmd = _default_windows_path
except ImportError:
    pass


def extract_text(path: str) -> str:
    """Return best-effort plain text for a single uploaded file."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        elif ext == ".docx":
            return _extract_docx(path)
        elif ext == ".doc":
            return _extract_doc_fallback(path)
        elif ext == ".txt":
            return _extract_txt(path)
        elif ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif"):
            return _extract_image_ocr(path)
        else:
            return ""
    except Exception as e:
        # Never let one bad file crash the whole analysis — surface the
        # failure as text so it's visible in the combined prompt/debug log
        # instead of silently vanishing.
        return f"[Could not read {os.path.basename(path)}: {e}]"


def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages).strip()
    if text:
        return text
    # No extractable text layer -> it's likely a scanned/image-only PDF.
    # Full PDF-to-image OCR needs poppler (pdf2image) which is an extra
    # system dependency, so we just flag it rather than silently returning
    # nothing.
    return "[PDF appears to be scanned/image-only; no extractable text found]"


def _extract_docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def _extract_doc_fallback(path: str) -> str:
    # Legacy .doc (pre-2007 binary format) isn't readable by python-docx.
    # Properly parsing it needs antiword/textract or LibreOffice headless
    # conversion. We surface that clearly instead of pretending it worked.
    return "[Legacy .doc format not parsed — please re-save as .docx or PDF]"


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read().strip()


def _extract_image_ocr(path: str) -> str:
    if _pytesseract is None:
        return "[Image OCR unavailable — Tesseract is not installed on this server. Please upload a PDF or DOCX instead.]"
    from PIL import Image
    img = Image.open(path)
    return _pytesseract.image_to_string(img, config="--oem 3 --psm 6").strip()


def extract_all(paths: list[str]) -> str:
    """
    Combine multiple uploaded files into one labeled text block, ready to
    drop straight into the OpenAI prompt as a single user message.
    """
    chunks = []
    for path in paths:
        filename = os.path.basename(path)
        text = extract_text(path)
        if not text:
            text = "[No text could be extracted from this file]"
        chunks.append(f"=== FILE: {filename} ===\n{text}")
    return "\n\n".join(chunks)
